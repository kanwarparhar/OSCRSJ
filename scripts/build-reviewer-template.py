#!/usr/bin/env python3
"""
Build the OSCRSJ Reviewer Template .docx at public/downloads/oscrsj-reviewer-template.docx.

Run from the repo root: `python3 scripts/build-reviewer-template.py`

Requires `python-docx` (`pip install python-docx`).

History:
  v1.0 (2026-04-25, Session 25) — original 7-section template generated via Node `docx` lib.
  v1.1 (2026-04-26, Session 33 follow-up) — Section 3 intro reframed as guidance scaffold,
        Section 3.4 rewritten to acknowledge that the structured review form has a single
        Feedback-and-review field (per Session 32 commit ea85525 + Session 33 commit
        1a56a85) so confidential concerns must be emailed separately, Section 7 rewritten
        for paste-into-textarea submission workflow. Build path migrated from Node docx-js
        (script preserved only in outputs scratchpad in Session 25) to Python python-docx
        committed in-tree so future sessions have a regenerator.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "public" / "downloads" / "oscrsj-reviewer-template.docx"

BROWN_DARK = RGBColor(0x3D, 0x2A, 0x18)
INK = RGBColor(0x12, 0x0D, 0x08)
TAN = RGBColor(0x99, 0x7E, 0x67)
BORDER = "997E67"


def set_cell_border(cell, color: str = BORDER, sz: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        tc_borders.append(e)
    tc_pr.append(tc_borders)


def set_paragraph_spacing(par, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6, line=1.15)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = BROWN_DARK
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def add_para(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    bold: bool = False,
    size: int = 10,
    after: int = 6,
) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = INK


def add_blank_lines(doc: Document, n: int = 3) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line=1.4)
        p.add_run(" ")


def add_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1.2)
    run = p.add_run(label)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = TAN


def build() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = INK

    # ===== Header =====
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSCRSJ")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BROWN_DARK

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Orthopedic Surgery Case Reports & Series Journal")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = BROWN_DARK

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REVIEWER TEMPLATE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TAN

    add_para(
        doc,
        "Recommended scaffold for organizing every peer review submitted to OSCRSJ. "
        "The actual review is submitted through the structured review form linked from "
        "your invitation email — this document is not uploaded.",
        italic=True,
        size=10,
        after=14,
    )

    # ===== Section 1 — Reviewer & Manuscript Information =====
    add_heading(doc, "1. Reviewer & Manuscript Information", level=1)

    fields = [
        ("Reviewer Name", "[Your full name]"),
        ("Reviewer Email", "[Your email]"),
        ("Affiliation", "[Your institution / department]"),
        ("ORCID iD", "[https://orcid.org/0000-0000-0000-0000]"),
        ("Manuscript ID", "[OSCRSJ-2026-####]"),
        ("Manuscript Title", "[Title as it appears in the invitation]"),
        ("Date of Review", "[YYYY-MM-DD]"),
        ("Conflict of Interest", "[None / Describe any actual or perceived conflict]"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(11.0)
    for i, (label, placeholder) in enumerate(fields):
        cells = table.rows[i].cells
        cells[0].width = Cm(5.0)
        cells[1].width = Cm(11.0)
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        for cell in cells[0].paragraphs:
            cell.paragraph_format.space_after = Pt(2)
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BROWN_DARK
        run = cells[1].paragraphs[0].add_run(placeholder)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = TAN

    add_blank_lines(doc, 1)

    # ===== Section 2 — AI Prohibition =====
    add_heading(doc, "2. AI Prohibition (Read Before Beginning)", level=1)
    add_para(
        doc,
        "The use of large language models, generative AI tools, or any automated review-writing "
        "assistance is strictly prohibited at OSCRSJ. This includes ChatGPT, Claude, Gemini, "
        "Perplexity, and any other AI system, whether used to generate reviewer comments, "
        "summarize the manuscript, draft the recommendation, or paraphrase your own notes.",
    )
    add_para(
        doc,
        "Uploading any portion of an unpublished manuscript to a third-party AI service is a "
        "breach of the confidentiality agreement you accepted when invited to review. Reviews "
        "suspected of AI generation will be discarded, the reviewer will be removed from the "
        "OSCRSJ reviewer pool, and the relevant institutional and professional bodies may be "
        "notified.",
    )
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10)
    run = p.add_run(
        "I confirm that this review was written by me, without the use of generative AI tools.    ☐  Yes"
    )
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = INK

    # ===== Section 3 — Structuring Your Review =====
    add_heading(doc, "3. Structuring Your Review", level=1)
    add_para(
        doc,
        "These blocks help you organize a complete review. Work through them in order — when "
        "you are ready to submit, paste the content of Sections 3.1, 3.2, 3.3, and 3.5 into the "
        "Feedback and review field on the structured review form (linked from your invitation "
        "email). Sections 4–6 are guidance, not required completions.",
        italic=True,
    )

    # 3.1 Summary
    add_heading(doc, "3.1 Summary", level=2)
    add_para(
        doc,
        "2–3 sentences. Brief synopsis of the manuscript and your overall impression of its "
        "contribution.",
    )
    add_label(doc, "Your summary:")
    add_blank_lines(doc, 4)

    # 3.2 Major Comments
    add_heading(doc, "3.2 Major Comments", level=2)
    add_para(
        doc,
        "Issues that must be addressed before the manuscript can be considered for publication. "
        "These might include missing clinical details, flawed reasoning, unsupported "
        "conclusions, or ethical concerns. Number each comment.",
    )
    add_label(doc, "Major comments:")
    add_blank_lines(doc, 8)

    # 3.3 Minor Comments
    add_heading(doc, "3.3 Minor Comments", level=2)
    add_para(
        doc,
        "Suggestions for improvement that would strengthen the manuscript but are not "
        "essential. These might include wording suggestions, additional references, or figure "
        "quality improvements. Number each comment.",
    )
    add_label(doc, "Minor comments:")
    add_blank_lines(doc, 6)

    # 3.4 Confidential concerns — REWRITTEN for v1.1
    add_heading(doc, "3.4 Concerns for the Editorial Office Only", level=2)
    add_para(
        doc,
        "The structured review form has a single Feedback and review field — everything you "
        "paste into it is visible to the authors after the editorial decision is communicated. "
        "There is no separate confidential channel in the form.",
    )
    add_para(
        doc,
        "If you have concerns intended ONLY for the editorial team — suspected ethics violations, "
        "plagiarism, undisclosed conflicts of interest, or questions about the editorial "
        "decision — DO NOT include them in the Feedback and review field. Email them separately "
        "to editorial@oscrsj.com referencing the Manuscript ID, and the editor will handle them "
        "outside the author-visible record.",
        bold=False,
    )
    add_label(doc, "Notes for editorial email (do NOT paste into the form):")
    add_blank_lines(doc, 4)

    # 3.5 Recommendation
    add_heading(doc, "3.5 Recommendation", level=2)
    add_para(doc, "Mark exactly one decision below. Mirror this choice in the form.")
    for label in (
        "☐  Accept — Manuscript is ready for publication as submitted.",
        "☐  Minor Revision — Small changes needed; does not require re-review.",
        "☐  Major Revision — Substantial changes needed; will require re-review.",
        "☐  Reject — Does not meet standards and cannot be improved through revision.",
    ):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=2)
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.color.rgb = INK

    add_label(doc, "Brief rationale for your recommendation:")
    add_blank_lines(doc, 4)

    # ===== Section 4 — Section-by-Section Evaluation Prompts =====
    add_heading(doc, "4. Section-by-Section Evaluation Prompts", level=1)
    add_para(
        doc,
        "Use these prompts as guides while you read. They are not required completions — answer "
        "where you have substantive feedback. Cite manuscript line numbers where possible.",
        italic=True,
    )

    prompts = [
        (
            "4.1 Title and Abstract",
            [
                "Is the title accurate, concise, and descriptive of the case?",
                "Does the abstract follow the required structure?",
                "Does the abstract stand alone as a complete summary?",
                "Are there abbreviations or references in the abstract? (There should not be.)",
            ],
        ),
        (
            "4.2 Introduction",
            [
                "Is the clinical context clearly established?",
                "Is the novelty or educational value of the case clearly stated?",
                "Is relevant background literature cited?",
                "Is it appropriately brief (not a full literature review)?",
            ],
        ),
        (
            "4.3 Case Presentation / Methods",
            [
                "Is the clinical timeline clear and logical?",
                "Are patient demographics, history, examination findings, and investigations adequately described?",
                "Is the surgical technique or treatment described in sufficient detail for reproduction?",
                "Are outcomes and follow-up duration clearly reported?",
                "Has all protected health information (PHI) been removed?",
            ],
        ),
        (
            "4.4 Discussion",
            [
                "Is the case placed in context of existing literature?",
                "Are differential diagnoses discussed?",
                "Is the rationale for the chosen treatment explained?",
                "Are limitations of the case report acknowledged?",
                "Are the conclusions supported by the evidence presented?",
            ],
        ),
        (
            "4.5 Figures and Tables",
            [
                "Are images of sufficient quality for publication?",
                "Are figure legends accurate and complete?",
                "Are patient identifiers removed from all images?",
                "Do figures and tables add value, or are any redundant?",
            ],
        ),
        (
            "4.6 References",
            [
                "Are references current and relevant?",
                "Are key foundational and recent references included?",
                "Are references formatted in Vancouver style?",
            ],
        ),
        (
            "4.7 Ethics and Reporting",
            [
                "Is a patient consent statement included?",
                "Is a conflict of interest statement included?",
                "Has the CARE checklist been completed (for case reports)?",
                "Has the JBI checklist been completed (for case series)?",
            ],
        ),
    ]
    for heading, bullets in prompts:
        add_heading(doc, heading, level=2)
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=2, line=1.2)
            run = p.add_run(b)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
        add_label(doc, "Notes:")
        add_blank_lines(doc, 3)

    # ===== Section 5 — Decision Categories =====
    add_heading(doc, "5. Decision Categories — Detailed Guidance", level=1)
    add_para(
        doc,
        "Use this guidance when selecting your recommendation in Section 3.5.",
        italic=True,
    )

    decisions = [
        (
            "Accept",
            "Use sparingly. Reserve for manuscripts that need no further revision and could be "
            "published in their current form. In practice, almost no first-round submission "
            "warrants Accept.",
        ),
        (
            "Minor Revision",
            "The manuscript is fundamentally sound. Issues are limited to clarifications, "
            "additional citations, figure quality, or minor wording. The revised manuscript will "
            "not need to return to peer review — the editor can confirm the changes.",
        ),
        (
            "Major Revision",
            "The manuscript has substantive issues that must be addressed before it can be "
            "considered for publication, but the underlying case and analysis are valuable. The "
            "revised manuscript will return to peer review (often to the same reviewer) for "
            "re-evaluation.",
        ),
        (
            "Reject",
            "The manuscript cannot be improved into a publishable form through revision. Common "
            "reasons: insufficient novelty, fatal methodological flaws, missing patient consent / "
            "IRB documentation, or scope misalignment. Provide constructive feedback even when "
            "rejecting — the authors may revise for another venue.",
        ),
    ]
    for label, desc in decisions:
        add_heading(doc, label, level=3)
        add_para(doc, desc)

    # ===== Section 6 — Guidance Tips =====
    add_heading(doc, "6. Guidance Tips for Writing a Great Review", level=1)

    tips = [
        (
            "Be specific",
            "Instead of “the discussion is weak,” say “the discussion does not address "
            "the alternative diagnosis of X, which presents similarly.” Cite manuscript line "
            "numbers wherever possible.",
        ),
        (
            "Be constructive",
            "Frame feedback as suggestions, not criticisms. “Consider adding a timeline "
            "figure” is better than “the chronology is confusing.”",
        ),
        (
            "Be respectful",
            "Authors put significant work into their submissions, often outside their clinical "
            "hours. Critique the work, never the author. Avoid sarcasm, condescension, or "
            "rhetorical questions implying incompetence.",
        ),
        (
            "Keep editorial-only concerns out of the form",
            "The Feedback and review field on the form is author-visible. Concerns intended only "
            "for the editorial team (Section 3.4) must be emailed separately to "
            "editorial@oscrsj.com — never paste them into the form.",
        ),
        (
            "Avoid rewriting",
            "Point out what needs improvement, but let the authors make the changes in their own "
            "voice. Do not paste rewritten paragraphs in your review — describe what is missing "
            "or unclear and let the authors decide how to address it.",
        ),
        (
            "Cite line numbers and figure numbers",
            "Every comment should include a precise location reference (e.g., “L142–148,” "
            "“Figure 2,” “Table 1, row 4”). Vague references waste author time and "
            "increase the chance the wrong passage gets revised.",
        ),
        (
            "Distinguish opinion from evidence",
            "When you disagree with the authors’ interpretation, say so as your view. When the "
            "authors have made a factual or statistical error, cite the source. “I would have "
            "approached this differently” is opinion; “Reference 12 contradicts this claim” "
            "is evidence.",
        ),
    ]
    for label, desc in tips:
        add_heading(doc, label, level=3)
        add_para(doc, desc)

    # ===== Section 7 — Submitting Your Review (REWRITTEN) =====
    add_heading(doc, "7. Submitting Your Review", level=1)
    add_para(
        doc,
        "Reviews are submitted through the structured review form on the OSCRSJ website — not as "
        "a Word document upload. Do NOT email this completed template to anyone.",
        italic=True,
    )

    submit_steps = [
        (
            "1.",
            "Open the structured review form using the link in your invitation email "
            "(URL pattern: oscrsj.com/review/[token]/form). You must accept the invitation "
            "before the form will open.",
        ),
        (
            "2.",
            "Copy the content of Sections 3.1 (Summary), 3.2 (Major Comments), 3.3 (Minor "
            "Comments), and your rationale from 3.5 into the Feedback and review field on the "
            "form. The form has a single freeform text field; structure within the field is up "
            "to you (numbered headings recommended).",
        ),
        (
            "3.",
            "Set the six rating scales on the form (Manuscript Quality, Novelty & Significance, "
            "Scientific Rigor & Methods, Data Quality & Results, Clarity & Presentation, and "
            "Compliance with Journal Scope), select your recommendation, and complete the "
            "conflict-of-interest declaration.",
        ),
        (
            "4.",
            "If you have concerns ONLY for the editorial team (Section 3.4), email them "
            "separately to editorial@oscrsj.com referencing the Manuscript ID. Do NOT paste them "
            "into the form.",
        ),
        (
            "5.",
            "Submit. You will receive a confirmation email and your contribution will be logged "
            "toward annual reviewer recognition.",
        ),
    ]
    for n, body in submit_steps:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line=1.3)
        run = p.add_run(f"{n}  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BROWN_DARK
        run = p.add_run(body)
        run.font.size = Pt(10)
        run.font.color.rgb = INK

    add_para(
        doc,
        "Reviews are due 21 days from your acceptance of the invitation. If you need an "
        "extension (up to 7 days), contact editorial@oscrsj.com BEFORE your deadline.",
        bold=True,
        after=14,
    )

    # ===== Footer =====
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSCRSJ Editorial Office  ·  editorial@oscrsj.com  ·  www.oscrsj.com")
    run.font.size = Pt(9)
    run.font.color.rgb = TAN

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reviewer Template v1.1  ·  Issued April 2026")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = TAN

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(REPO_ROOT)} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
