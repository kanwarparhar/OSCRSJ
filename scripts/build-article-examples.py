#!/usr/bin/env python3
"""
Build all 6 OSCRSJ worked-example .docx files at
public/downloads/oscrsj-example-{type}.docx.

Run from the repo root: `python3 scripts/build-article-examples.py`

Requires `python-docx` (`pip install python-docx`).

This script regenerates the per-article-type WORKED EXAMPLES (the populated
versions of the templates that authors download to see what good looks like).
Every example is re-rendered under the same 8 v1.1 formatting locks used by
scripts/build-article-templates.py:

  1. All headings render in sentence case (only the first letter capitalised),
     bold, no italics.
  2. The Introduction (or first body section) starts on a new page after the
     abstract + keywords block.
  3. Every heading carries an empty paragraph ABOVE it, never below.
  4. The document carries a header with a short running title in the top-right
     corner of every page.
  5. In-text citations render as superscript Word hyperlinks anchored to
     bookmarks (REF_1, REF_2, …) seeded at the matching reference list entry.
  6. References sits on its own page with a centred sentence-case heading and
     no blank line between the heading and the first reference.
  7. Figure legends mirrors References — its own page, centred sentence-case
     heading, no blank line between heading and first legend.
  8. Body text is Times New Roman 12 pt, double-spaced, all black.

Citation markup convention used in the section-builder functions:

  Use the helper `cite(*nums)` to create a citation segment. For example,
  `cite(1)`         renders as a single superscript hyperlink to REF_1.
  `cite(2, 3)`      renders as `2,3` (comma-separated) hyperlinks to REF_2
                    and REF_3.

  Mix plain prose with citations using `runs(...)`:
      runs("This finding is consistent with prior work", cite(1, 2),
           ", and a recent meta-analysis confirmed the trend",
           cite(3), ".")

History:
  v1.0 (2026-05-04) — initial release. Replaces the 2026-04-25 worked-example
                       set that pre-dated the v1.1 template formatting locks.
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths / fonts / colour constants
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
DOWNLOADS_DIR = REPO_ROOT / "public" / "downloads"

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
LINE_SPACING = 2.0  # double-spaced
BLACK = RGBColor(0x00, 0x00, 0x00)
HYPERLINK_BLUE_HEX = "0563C1"


# ---------------------------------------------------------------------------
# Low-level docx helpers (kept byte-compatible with build-article-templates.py)
# ---------------------------------------------------------------------------


def _set_default_run_font(run) -> None:
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE_PT)
    run.font.color.rgb = BLACK


def _apply_paragraph_defaults(par, *, double_spaced: bool = True) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if double_spaced:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_blank_paragraph(doc) -> None:
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p)


def add_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: Optional[int] = None,
    indent_hanging_inches: Optional[float] = None,
    double_spaced: bool = True,
) -> object:
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, double_spaced=double_spaced)
    if align is not None:
        p.alignment = align
    if indent_hanging_inches is not None:
        pf = p.paragraph_format
        pf.left_indent = Pt(36)
        pf.first_line_indent = Pt(-36)
    run = p.add_run(text)
    _set_default_run_font(run)
    run.bold = bold
    run.italic = italic
    return p


def add_heading_paragraph(
    doc,
    text: str,
    *,
    centered: bool = False,
    page_break_before: bool = False,
) -> object:
    """Bold, sentence-case, no italics. Optionally centred, optionally on a
    new page."""
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if page_break_before:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    _set_default_run_font(run)
    run.bold = True
    run.italic = False
    return p


def add_bookmark(paragraph, name: str) -> None:
    """Insert a Word bookmark wrapping the START of the paragraph's first run.

    Bookmark IDs are kept globally unique by walking the document's existing
    bookmark IDs at call time."""
    body = paragraph._parent._element
    existing_ids = [
        int(b.get(qn("w:id")))
        for b in body.iter(qn("w:bookmarkStart"))
        if b.get(qn("w:id")) is not None
    ]
    next_id = (max(existing_ids) + 1) if existing_ids else 0

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(next_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(next_id))

    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(end)
        pPr.addnext(start)
    else:
        p.insert(0, end)
        p.insert(0, start)


def add_internal_hyperlink_run(
    paragraph,
    text: str,
    bookmark_name: str,
    *,
    superscript: bool = True,
) -> None:
    """Append a run linked to an internal bookmark; styled blue + superscript."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rPr.append(rfonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), HYPERLINK_BLUE_HEX)
    rPr.append(color)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(FONT_SIZE_PT * 2))
    rPr.append(sz)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if superscript:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        rPr.append(vert)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Citation markup — `runs(...)` builds a single paragraph from a sequence of
# plain strings + Citation tokens. Citation tokens render as one or more
# superscript hyperlinks separated by commas (no spaces).
# ---------------------------------------------------------------------------


class Citation:
    """A token representing one or more reference numbers that should render
    as comma-separated superscript hyperlinks (e.g., 2,3 → REF_2 + REF_3)."""

    __slots__ = ("nums",)

    def __init__(self, nums: Tuple[int, ...]) -> None:
        if not nums:
            raise ValueError("Citation must wrap at least one reference number.")
        self.nums = nums


def cite(*nums: int) -> Citation:
    return Citation(nums)


Segment = Union[str, Citation]


def add_paragraph_with_citations(
    doc,
    segments: List[Segment],
    *,
    double_spaced: bool = True,
) -> object:
    """Append a single paragraph by concatenating string segments and
    superscript-hyperlink Citation segments."""
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, double_spaced=double_spaced)
    for seg in segments:
        if isinstance(seg, Citation):
            for i, n in enumerate(seg.nums):
                if i > 0:
                    # Comma separator BETWEEN superscripts. Keep the comma
                    # itself superscript so the whole citation cluster reads
                    # as one tight superscript block.
                    add_internal_hyperlink_run(p, ",", f"REF_{seg.nums[0]}",
                                               superscript=True)
                add_internal_hyperlink_run(p, str(n), f"REF_{n}",
                                           superscript=True)
        else:
            run = p.add_run(seg)
            _set_default_run_font(run)
    return p


def runs(*segments: Segment) -> List[Segment]:
    """Sugar so callers can write `runs("text", cite(1), ".")` instead of
    `[..]`."""
    return list(segments)


# ---------------------------------------------------------------------------
# Section-wide plumbing: header, line numbering, margins
# ---------------------------------------------------------------------------


def configure_section(section, *, running_title: str) -> None:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sect_pr.append(ln)

    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(running_title)
    _set_default_run_font(run)
    run.italic = True


def configure_document_defaults(doc) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


# ---------------------------------------------------------------------------
# Section builder — section dicts mirror build-article-templates.py kinds:
#   {"kind": "h1", "text": "Heading", "page_break_before": False}
#   {"kind": "h2", "text": "Subheading"}
#   {"kind": "para", "text": "Plain prose."}
#   {"kind": "para_runs", "segments": [...]}   ← citations + plain text
# ---------------------------------------------------------------------------


def main_heading(text: str, *, page_break_before: bool = False) -> dict:
    return {"kind": "h1", "text": text, "page_break_before": page_break_before}


def sub_heading(text: str) -> dict:
    return {"kind": "h2", "text": text}


def para(text: str) -> dict:
    return {"kind": "para", "text": text}


def para_runs(*segments: Segment) -> dict:
    return {"kind": "para_runs", "segments": list(segments)}


# ---------------------------------------------------------------------------
# Worked-example content — one builder per article type.
#
# Content was migrated verbatim from the 2026-04-25 worked-example set, then
# re-marked-up: every inline `[1]`, `[2,3]`, etc. citation was replaced with a
# `cite(1)` / `cite(2, 3)` token so it renders as a superscript hyperlink to
# the matching REF_n bookmark in the Reference list.
# ---------------------------------------------------------------------------


def case_report_example_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        sub_heading("Introduction"),
        para(
            "Iatrogenic median nerve injury after closed reduction and "
            "percutaneous pinning of pediatric supracondylar humerus "
            "fractures is uncommon and has been reported predominantly in "
            "association with medial-entry K-wires. Isolated lateral-entry-"
            "wire-mediated median nerve injury is rarer still."
        ),
        sub_heading("Case presentation"),
        para(
            "A 7-year-old previously healthy boy presented with a Gartland "
            "type III supracondylar humerus fracture without preoperative "
            "neurovascular deficit. Closed reduction and percutaneous "
            "pinning with three lateral-entry K-wires was performed; the "
            "postoperative bedside examination on day 1 demonstrated a "
            "complete sensorimotor median nerve palsy. High-resolution "
            "ultrasound localized the deficit to a kinked median nerve "
            "coursing across the most medial of the lateral-entry K-wires. "
            "The patient was returned to the operating theater on day 2 for "
            "wire removal, nerve decompression, and revision pinning. Full "
            "motor and sensory recovery was achieved by 16 weeks "
            "postoperatively."
        ),
        sub_heading("Discussion"),
        para(
            "This case illustrates that iatrogenic median nerve injury can "
            "occur with lateral-entry K-wires when the most medial of "
            "multiple lateral wires is advanced steeply across the medial "
            "column. Repeated postoperative neurovascular examination — "
            "including a documented exam in the post-anesthesia care unit "
            "— is essential. High-resolution ultrasound was the most "
            "informative imaging modality for localizing the lesion "
            "preoperatively."
        ),
        sub_heading("Conclusion"),
        para(
            "Iatrogenic median nerve injury after pediatric supracondylar "
            "pinning is uncommon, can occur with lateral-entry as well as "
            "medial-entry K-wires, and is recoverable when recognized early."
        ),
        main_heading("Keywords"),
        para(
            "supracondylar humerus fracture; median nerve; iatrogenic "
            "injury; pediatric orthopedics; percutaneous pinning"
        ),
        main_heading("Introduction", page_break_before=True),
        para_runs(
            "Supracondylar humerus fractures are the most common elbow "
            "fracture in children, accounting for approximately 60–70% of "
            "all pediatric elbow fractures and most often occurring between "
            "the ages of five and seven years",
            cite(1),
            ". Closed reduction and percutaneous pinning is the standard of "
            "care for displaced (Gartland type II and III) extension-type "
            "fractures and yields excellent radiographic and functional "
            "outcomes when performed with attention to wire trajectory and "
            "stability",
            cite(2, 3),
            ".",
        ),
        para_runs(
            "Iatrogenic neurovascular injury is a recognized but uncommon "
            "complication. The reported incidence of post-pinning ulnar "
            "nerve injury attributable to medial-entry K-wires ranges from "
            "1% to 5% in modern series",
            cite(4),
            ". Median nerve injury is rarer still and has been reported "
            "predominantly in association with anterior interosseous branch "
            "involvement after preoperative neurologic compromise; "
            "iatrogenic isolated median nerve injury attributable to a "
            "percutaneous K-wire has been described only in case reports",
            cite(5),
            ".",
        ),
        para(
            "We describe a case of complete iatrogenic median nerve palsy "
            "in a 7-year-old boy following closed reduction and "
            "percutaneous pinning of a Gartland type III supracondylar "
            "humerus fracture, with full motor and sensory recovery after "
            "wire removal and nerve decompression."
        ),
        main_heading("Case presentation"),
        para(
            "A 7-year-old previously healthy right-hand-dominant boy "
            "presented to the emergency department after a fall from monkey "
            "bars onto an outstretched right arm, with immediate pain and "
            "inability to use the right elbow. Initial neurovascular "
            "examination demonstrated intact radial, median, and ulnar "
            "nerve function with palpable radial pulse and well-perfused "
            "fingertips. Anteroposterior and lateral radiographs "
            "demonstrated a Gartland type III extension-type supracondylar "
            "humerus fracture with complete posteromedial displacement of "
            "the distal fragment (Figure 1). No associated upper-extremity "
            "injury was identified on a complete trauma series."
        ),
        para(
            "The patient was taken to the operating theater within four "
            "hours of presentation. Closed reduction was achieved on the "
            "first attempt under fluoroscopic guidance, and the fracture "
            "was stabilized with three lateral-entry K-wires (two divergent "
            "lateral, one capitellar). The construct was assessed for "
            "stability under image intensification and the elbow was "
            "immobilized in a long-arm posterior splint at 70° of flexion. "
            "Estimated blood loss was less than 20 mL; tourniquet time was "
            "32 minutes."
        ),
        para(
            "On postoperative day 1, the patient was noted to have a new "
            "complete sensorimotor median nerve deficit in the right hand. "
            "High-resolution ultrasound performed the same day localized "
            "the deficit to a kinked median nerve passing across the most "
            "medial of the three lateral-entry K-wires. The patient was "
            "returned to the operating theater on postoperative day 2 for "
            "wire removal, nerve decompression, and revision pinning. The "
            "clinical course is summarized in Table 1."
        ),
        para(
            "The patient was permitted toe-touch weight-bearing for six "
            "weeks postoperatively, progressing to full weight-bearing by "
            "ten weeks. He completed an outpatient physical therapy "
            "program. At 16-week follow-up, he had full active range of "
            "motion of the elbow and a fully recovered median nerve "
            "examination, and he had returned to age-appropriate activity. "
            "Surveillance examination at 12 months showed no residual "
            "deficit."
        ),
        para(
            "Written informed consent for publication of this case report "
            "and accompanying images was obtained from the parent and legal "
            "guardian of the minor patient, in accordance with OSCRSJ "
            "patient consent policy (see oscrsj.com/templates)."
        ),
        main_heading("Discussion"),
        para_runs(
            "Iatrogenic median nerve injury after closed reduction and "
            "percutaneous pinning of pediatric supracondylar humerus "
            "fractures is rare. In the largest single-institution series "
            "to date, Brown and colleagues identified six iatrogenic nerve "
            "injuries among 1,242 pinned supracondylar fractures (0.5%); "
            "two of these involved the median nerve, and both were "
            "associated with medial-entry K-wires",
            cite(6),
            ". Our case is unusual in that the offending wire was placed "
            "laterally — highlighting that median nerve compromise is not "
            "exclusive to medial wire placement, particularly when the "
            "most medial of multiple lateral-entry wires is advanced "
            "steeply across the medial column.",
        ),
        para_runs(
            "Two findings from our case have practical relevance for the "
            "operating surgeon. First, repeated postoperative "
            "neurovascular examination — including a documented "
            "examination performed in the recovery room before the patient "
            "is discharged from the post-anesthesia care unit — is "
            "essential. The deficit in our patient was not present at the "
            "conclusion of surgery on the in-room exam, suggesting that "
            "nerve compromise developed either during transfer or during "
            "the immediate postoperative period as edema increased around "
            "the fracture site. Second, high-resolution ultrasound was the "
            "single most informative imaging modality for localizing the "
            "lesion preoperatively",
            cite(7),
            ". Magnetic resonance imaging would have required general "
            "anesthesia in this age group and would have delayed revision "
            "surgery; ultrasound was performed at the bedside within 90 "
            "minutes of identifying the deficit.",
        ),
        para(
            "The principal limitation of any case report is the absence of "
            "a comparison population. We cannot infer the rate at which "
            "similar wire trajectories produce nerve injury, nor can we "
            "conclude that earlier ultrasound shortened the time to "
            "recovery. Reporting this case nevertheless adds to the small "
            "literature describing isolated lateral-wire-mediated median "
            "nerve injury and offers a practical example of bedside "
            "ultrasound localization in pediatric iatrogenic nerve "
            "compromise."
        ),
        main_heading("Conclusion"),
        para(
            "Iatrogenic median nerve injury after closed reduction and "
            "percutaneous pinning of pediatric supracondylar humerus "
            "fractures is uncommon, can occur with lateral-entry as well "
            "as medial-entry K-wires, and is recoverable when recognized "
            "early. Repeated postoperative neurovascular examination — "
            "including a documented exam in the post-anesthesia care unit "
            "— should be standard, and high-resolution ultrasound is a "
            "useful first-line imaging modality when wire-related nerve "
            "injury is suspected."
        ),
    ]


CASE_REPORT_REFERENCES = [
    "Skaggs DL, Cluck MW, Mostofi A, Flynn JM, Kay RM. Lateral-entry pin fixation in the management of supracondylar fractures in children. J Bone Joint Surg Am. 2004;86(4):702–707. https://doi.org/10.2106/00004623-200404000-00006",
    "Bahk MS, Srikumaran U, Ain MC, et al. Patterns of pediatric supracondylar humerus fractures. J Pediatr Orthop. 2008;28(5):493–499. https://doi.org/10.1097/BPO.0b013e31817bb860",
    "Babal JC, Mehlman CT, Klein G. Nerve injuries associated with pediatric supracondylar humeral fractures: a meta-analysis. J Pediatr Orthop. 2010;30(3):253–263. https://doi.org/10.1097/BPO.0b013e3181d213a6",
    "Lyons ST, Quinn M, Stanitski CL. Neurovascular injuries in type III humeral supracondylar fractures in children. Clin Orthop Relat Res. 2000;376:62–67. https://doi.org/10.1097/00003086-200007000-00010",
    "Mubarak SJ, Davids JR. Closed treatment of supracondylar fractures of the humerus in children. Tech Orthop. 1989;4(1):51–58.",
    "Brown IC, Zinar DM. Traumatic and iatrogenic neurological complications after supracondylar humerus fractures in children. J Pediatr Orthop. 1995;15(4):440–443. https://doi.org/10.1097/01241398-199507000-00005",
    "Bargagliotti M, Pomero V, Maes-Clavier C, et al. The role of ultrasound in pediatric peripheral nerve injuries. J Ultrasound Med. 2018;37(11):2557–2569. https://doi.org/10.1002/jum.14609",
]


CASE_REPORT_FIGURES = [
    "Figure 1.  Pre-operative anteroposterior (A) and lateral (B) radiographs of the right elbow demonstrating a Gartland type III extension-type supracondylar humerus fracture with complete posteromedial displacement of the distal fragment.",
    "Figure 2.  Intraoperative fluoroscopic image after revision pinning showing two crossed lateral-entry K-wires in stable configuration.",
]


def case_series_example_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        sub_heading("Background"),
        para(
            "Femoral head osteonecrosis after high-dose corticosteroid "
            "therapy for severe COVID-19 pneumonia is increasingly "
            "recognized but its incidence in patients receiving cumulative "
            "doses below the historical 2,000 mg prednisone-equivalent "
            "threshold is not well characterized."
        ),
        sub_heading("Methods"),
        para(
            "We retrospectively reviewed eight consecutive patients (12 "
            "hips) treated with staged bilateral or unilateral core "
            "decompression with autologous bone marrow aspirate concentrate "
            "(BMAC) at a single academic center between January 2022 and "
            "December 2024. Inclusion required Ficat stage I or II "
            "osteonecrosis on magnetic resonance imaging and prior "
            "dexamethasone exposure for severe COVID-19. Outcomes included "
            "Harris Hip Score, conversion to arthroplasty, and radiographic "
            "progression to subchondral collapse at minimum 18-month "
            "follow-up."
        ),
        sub_heading("Results"),
        para(
            "Mean age was 47 years (range 38–58); mean cumulative "
            "prednisone-equivalent exposure was 1,485 mg. At minimum "
            "18-month follow-up, 11 of 12 hips (92%) showed radiographic "
            "stability without progression to subchondral collapse. Mean "
            "Harris Hip Score improved from 53 preoperatively to 89 at "
            "final follow-up. One hip (1/12, 8%) progressed and was "
            "converted to total hip arthroplasty at 22 months."
        ),
        sub_heading("Discussion"),
        para(
            "Joint-preserving surgery with core decompression and BMAC "
            "injection yielded durable clinical improvement in 92% of pre-"
            "collapse hips after sub-threshold corticosteroid exposure. "
            "The single conversion occurred in the patient with the "
            "largest pre-operative lesion volume."
        ),
        sub_heading("Conclusion"),
        para(
            "Core decompression with BMAC may forestall arthroplasty in "
            "pre-collapse femoral head osteonecrosis after sub-threshold "
            "COVID-era corticosteroid exposure when initiated early in the "
            "disease course."
        ),
        main_heading("Keywords"),
        para(
            "avascular necrosis; femoral head; COVID-19; corticosteroids; "
            "core decompression; bone marrow aspirate concentrate"
        ),
        main_heading("Introduction", page_break_before=True),
        para_runs(
            "Avascular necrosis of the femoral head is a debilitating "
            "condition in which interruption of the subchondral blood "
            "supply produces ischemic death of marrow elements and "
            "trabecular bone, ultimately leading to subchondral collapse, "
            "joint incongruity, and end-stage osteoarthritis. More than "
            "80% of femoral heads with established osteonecrosis progress "
            "to collapse within five years and the majority require total "
            "hip arthroplasty in the third or fourth decade of life",
            cite(1),
            ".",
        ),
        para_runs(
            "The COVID-19 pandemic introduced a new at-risk population. "
            "Severe SARS-CoV-2 pneumonia was treated with extended courses "
            "of high-dose dexamethasone following the publication of the "
            "RECOVERY trial in mid-2020",
            cite(2),
            ". Although the prevailing dogma had previously held that the "
            "threshold for steroid-induced AVN required cumulative "
            "prednisone-equivalent exposure above 2,000 mg, recent reports "
            "describe AVN at substantially lower exposures in the setting "
            "of severe systemic inflammation",
            cite(3, 4),
            ".",
        ),
        para_runs(
            "Core decompression is the most extensively studied joint-"
            "preserving intervention for pre-collapse osteonecrosis, with "
            "reported success rates of 60% to 80% in Ficat stage I and II "
            "disease",
            cite(5),
            ". The addition of biologic augmentation, including autologous "
            "bone marrow aspirate concentrate, has been associated with "
            "improved outcomes in several recent meta-analyses",
            cite(6),
            ". We describe a single-center experience treating eight "
            "patients with sub-threshold corticosteroid exposure for "
            "COVID-19 pneumonia using core decompression with BMAC "
            "injection.",
        ),
        main_heading("Methods"),
        para(
            "We retrospectively reviewed all patients treated at a single "
            "academic medical center between January 2022 and December "
            "2024 for COVID-19-associated femoral head osteonecrosis with "
            "core decompression and autologous bone marrow aspirate "
            "concentrate injection. Inclusion criteria were: (1) Ficat "
            "stage I or II osteonecrosis on magnetic resonance imaging; "
            "(2) prior dexamethasone exposure for severe COVID-19 "
            "pneumonia; (3) cumulative prednisone-equivalent exposure "
            "between 500 and 2,000 mg; (4) age 18 years or older; (5) "
            "absence of alternative AVN etiology on hematologic workup. "
            "Exclusion criteria were prior hip surgery, sickle cell "
            "disease, alcoholism, and inadequate follow-up (less than 18 "
            "months)."
        ),
        para(
            "All patients underwent the same operative protocol. Each "
            "operation was performed under spinal anesthesia in supine "
            "position on a fluoroscopy-compatible table. A 2 cm lateral "
            "incision was made distal to the greater trochanter and three "
            "core decompression channels of 6 mm diameter each were "
            "created under fluoroscopic guidance. Sixty milliliters of "
            "bone marrow aspirate concentrate were prepared from the "
            "ipsilateral iliac crest using a commercial centrifugation "
            "system and delivered through the trephine track. For "
            "bilateral cases, the contralateral hip was treated 12 weeks "
            "after the index procedure."
        ),
        para(
            "Primary outcome was conversion to total hip arthroplasty by "
            "final follow-up. Secondary outcomes included Harris Hip "
            "Score, radiographic progression on serial anteroposterior and "
            "frog-leg lateral radiographs scored by a single orthopedic "
            "radiologist blinded to clinical outcome, and patient-reported "
            "return to activity. Institutional Review Board approval was "
            "obtained as a single-center retrospective chart review (see "
            "oscrsj.com/templates for IRB approval policy)."
        ),
        main_heading("Results"),
        para(
            "Eight patients (12 hips) met inclusion criteria during the "
            "study window. Patient demographics, cumulative steroid "
            "exposure, and outcomes are summarized in Table 1. Mean age "
            "was 47 years (range 38–58), and mean follow-up was 22 months "
            "(range 18–34)."
        ),
        para(
            "At minimum 18-month follow-up, 11 of 12 hips (92%) showed "
            "radiographic stability without progression to subchondral "
            "collapse. Mean Harris Hip Score improved from 53 "
            "preoperatively to 89 at final follow-up. Seven of eight "
            "patients (88%) reported return to recreational physical "
            "activity by 14 months postoperatively."
        ),
        para(
            "One hip (1/12, 8%) progressed and was converted to total hip "
            "arthroplasty at 22 months. This patient had the largest "
            "preoperative lesion volume in the cohort (47% of femoral "
            "head). Two minor complications occurred: a superficial wound "
            "dehiscence treated with local wound care and a transient "
            "lateral femoral cutaneous nerve neuropraxia that resolved by "
            "six weeks postoperatively. There were no infections, no "
            "episodes of femoral head fracture, and no thromboembolic "
            "events."
        ),
        main_heading("Discussion"),
        para_runs(
            "This single-center series describes the early outcomes of "
            "staged core decompression with bone marrow aspirate "
            "concentrate in 12 hips of patients with COVID-era "
            "corticosteroid-associated femoral head osteonecrosis treated "
            "below the historical threshold. The 92% radiographic "
            "stability rate at minimum 18-month follow-up is consistent "
            "with prior literature reporting 60–80% success rates for core "
            "decompression in Ficat stage I and II disease, and may "
            "reflect improved patient selection in a population in which "
            "the inciting exposure (corticosteroids) was time-limited and "
            "discontinued well before surgery",
            cite(5, 6),
            ".",
        ),
        para_runs(
            "The single conversion to arthroplasty occurred in the patient "
            "with the largest preoperative lesion volume (47%), consistent "
            "with prior literature identifying lesion size as a predictor "
            "of progression after core decompression",
            cite(7),
            ". This finding underscores the importance of magnetic "
            "resonance imaging in preoperative planning and patient "
            "counseling.",
        ),
        para(
            "Limitations include the small cohort size, single-center "
            "retrospective design, absence of a comparison group, and "
            "short follow-up duration. We cannot infer whether the "
            "addition of BMAC contributes to outcomes beyond what would be "
            "expected with core decompression alone in this population. A "
            "multicenter prospective registry would be the natural next "
            "step."
        ),
        main_heading("Conclusion"),
        para(
            "Core decompression with autologous bone marrow aspirate "
            "concentrate yielded durable clinical and radiographic "
            "outcomes in 92% of pre-collapse hips at minimum 18-month "
            "follow-up after sub-threshold COVID-era corticosteroid "
            "exposure. Joint-preserving surgery may forestall arthroplasty "
            "in this population when initiated early in the disease "
            "course. Larger prospective registries are needed."
        ),
    ]


CASE_SERIES_REFERENCES = [
    "Mont MA, Cherian JJ, Sierra RJ, Jones LC, Lieberman JR. Nontraumatic osteonecrosis of the femoral head: where do we stand today? J Bone Joint Surg Am. 2015;97(19):1604–1627. https://doi.org/10.2106/JBJS.O.00071",
    "RECOVERY Collaborative Group. Dexamethasone in hospitalized patients with COVID-19. N Engl J Med. 2021;384(8):693–704. https://doi.org/10.1056/NEJMoa2021436",
    "Zhao R, Wang H, Wang X, Feng F. Steroid therapy and the risk of osteonecrosis in SARS-CoV-2 patients: a systematic review. Osteoporos Int. 2022;33(6):1175–1185. https://doi.org/10.1007/s00198-021-06264-z",
    "Agarwala SR, Vijayvargiya M, Pandey P. Avascular necrosis as a part of long COVID-19. BMJ Case Rep. 2021;14(7):e242101. https://doi.org/10.1136/bcr-2021-242101",
    "Marker DR, Seyler TM, Ulrich SD, Srivastava S, Mont MA. Do modern techniques improve core decompression outcomes for hip osteonecrosis? Clin Orthop Relat Res. 2008;466(5):1093–1103. https://doi.org/10.1007/s11999-008-0184-9",
    "Hernigou P, Beaujean F. Treatment of osteonecrosis with autologous bone marrow grafting. Clin Orthop Relat Res. 2002;405:14–23. https://doi.org/10.1097/00003086-200212000-00003",
    "Steinberg ME, Hayken GD, Steinberg DR. A quantitative system for staging avascular necrosis. J Bone Joint Surg Br. 1995;77(1):34–41.",
]


CASE_SERIES_FIGURES = [
    "Figure 1.  Patient flow diagram for the study cohort across the January 2022–December 2024 inclusion window.",
    "Figure 2.  Pre- and postoperative coronal T1-weighted magnetic resonance images of a representative patient demonstrating stable lesion appearance at 18-month follow-up.",
]


def surgical_technique_example_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        para(
            "Schatzker type IV tibial plateau fractures involve the medial "
            "condyle and are commonly approached through an anteromedial "
            "incision, which provides limited visualization of the "
            "posteromedial fragment. We describe a modified posteromedial "
            "approach using a curved skin incision posterior to the medial "
            "collateral ligament, dissecting between the medial "
            "gastrocnemius and pes anserinus to expose the posteromedial "
            "fragment directly. This approach permits anatomic reduction "
            "and buttress plate fixation under direct visualization while "
            "sparing the medial collateral ligament. We illustrate the "
            "technique with intraoperative photographs and present early "
            "outcomes in a small series."
        ),
        main_heading("Keywords"),
        para(
            "tibial plateau fracture; Schatzker type IV; posteromedial "
            "approach; buttress plate; surgical technique"
        ),
        main_heading("Introduction", page_break_before=True),
        para_runs(
            "Schatzker type IV tibial plateau fractures involve the medial "
            "tibial condyle, often with a posteromedial fragment that is "
            "difficult to visualize through a standard anteromedial "
            "approach",
            cite(1, 2),
            ". Inadequate reduction of the posteromedial fragment has been "
            "associated with progressive varus collapse and early "
            "post-traumatic arthrosis",
            cite(3),
            ". Several posterior and posteromedial approaches have been "
            "described, but each carries trade-offs in terms of "
            "neurovascular risk, soft-tissue exposure, and the surgeon's "
            "familiarity with the operative window.",
        ),
        para(
            "We describe a modified posteromedial approach that uses a "
            "curved skin incision posterior to the medial collateral "
            "ligament, dissecting between the medial gastrocnemius and pes "
            "anserinus to expose the posteromedial fragment directly. The "
            "approach permits anatomic reduction and direct buttress plate "
            "fixation under direct visualization while sparing the medial "
            "collateral ligament."
        ),
        main_heading("Surgical technique"),
        para(
            "Step 1. Positioning. The patient is positioned supine on a "
            "radiolucent table with a small bump under the ipsilateral hip "
            "to allow approximately 30° of external rotation of the "
            "affected limb. A non-sterile tourniquet is applied to the "
            "proximal thigh and inflated to 300 mmHg after exsanguination."
        ),
        para(
            "Step 2. Skin incision. A curved 10–12 cm skin incision is "
            "made beginning 4 cm proximal to the medial joint line, "
            "curving posteriorly behind the medial collateral ligament, "
            "and extending distally along the posteromedial border of the "
            "tibia. The incision is centered over the posteromedial "
            "fragment as identified on preoperative computed tomography."
        ),
        para(
            "Step 3. Superficial dissection. The saphenous nerve and vein "
            "are identified at the proximal extent of the incision and "
            "protected with a vessel loop. Subcutaneous flaps are "
            "mobilized to expose the underlying fascial layer (Figure 1)."
        ),
        para(
            "Step 4. Deep interval. The interval between the medial head "
            "of the gastrocnemius posteriorly and the pes anserinus "
            "anteriorly is developed by sharp dissection. The medial "
            "gastrocnemius is retracted posteriorly with a Hohmann "
            "retractor, exposing the posteromedial capsule."
        ),
        para(
            "Step 5. Capsulotomy. A longitudinal capsulotomy is made "
            "directly over the posteromedial fragment under direct "
            "visualization. The fragment is mobilized with a small "
            "periosteal elevator (Figure 2)."
        ),
        para(
            "Step 6. Reduction and provisional fixation. Anatomic "
            "reduction is achieved by combined manual traction and direct "
            "manipulation of the posteromedial fragment. Provisional "
            "fixation is achieved with two 1.6 mm Kirschner wires placed "
            "from the posteromedial cortex into the lateral condyle."
        ),
        para(
            "Step 7. Definitive fixation. A 3.5 mm pre-contoured "
            "posteromedial buttress plate is placed flush against the "
            "posteromedial cortex. The plate is fixed with three or four "
            "3.5 mm cortical screws distally and two or three 3.5 mm "
            "cortical screws proximally that engage the subchondral bone "
            "of the posteromedial fragment under fluoroscopic guidance."
        ),
        para(
            "Step 8. Closure. The capsule is closed with 0 absorbable "
            "suture. The deep fascia is approximated with 2-0 absorbable "
            "suture. The skin is closed in standard layered fashion. A "
            "non-sterile drain is not routinely placed."
        ),
        para(
            "Step 9. Postoperative protocol. The patient is permitted "
            "touch-down weight-bearing with crutches for six weeks. Active "
            "and passive range-of-motion exercises are initiated on "
            "postoperative day 1. Progression to partial and full weight-"
            "bearing is permitted between weeks 6 and 10 based on "
            "radiographic evidence of healing."
        ),
        main_heading("Discussion"),
        para_runs(
            "The principal advantage of this modified posteromedial "
            "approach is direct visualization of the posteromedial "
            "fragment without disruption of the medial collateral "
            "ligament, which is a known concern with more anterior "
            "approaches",
            cite(4),
            ". The interval between the medial gastrocnemius and pes "
            "anserinus provides a safe operative window distant from the "
            "popliteal neurovascular structures and avoids the morbidity "
            "of the more posterior approaches described by previous "
            "authors",
            cite(5),
            ".",
        ),
        para(
            "The principal limitation of the approach is its visualization "
            "of the lateral condyle, which is poor; if the fracture "
            "pattern includes a lateral split or depression component, a "
            "separate anterolateral incision may be required and should be "
            "planned at the time of preoperative templating. The learning "
            "curve for the interval is short for surgeons familiar with "
            "the medial soft-tissue anatomy of the knee."
        ),
        main_heading("Conclusion"),
        para(
            "The modified posteromedial approach provides direct, well-"
            "illuminated access to the posteromedial fragment in Schatzker "
            "type IV tibial plateau fractures and permits anatomic "
            "reduction and buttress plate fixation under direct "
            "visualization. It should be considered when the fracture "
            "pattern is dominated by the posteromedial fragment and when "
            "MCL preservation is a priority."
        ),
    ]


SURGICAL_TECHNIQUE_REFERENCES = [
    "Schatzker J, McBroom R, Bruce D. The tibial plateau fracture. The Toronto experience 1968–1975. Clin Orthop Relat Res. 1979;138:94–104.",
    "Markhardt BK, Gross JM, Monu JU. Schatzker classification of tibial plateau fractures: use of CT and MR imaging improves assessment. Radiographics. 2009;29(2):585–597. https://doi.org/10.1148/rg.292085078",
    "Barei DP, Nork SE, Mills WJ, Coles CP, Henley MB, Benirschke SK. Functional outcomes of severe bicondylar tibial plateau fractures treated with dual incisions and medial and lateral plates. J Bone Joint Surg Am. 2006;88(8):1713–1721. https://doi.org/10.2106/JBJS.E.00907",
    "Galla M, Lobenhoffer P. The direct, dorsal approach to the treatment of unstable tibial posteromedial fracture-dislocations. Unfallchirurg. 2003;106(3):241–247. https://doi.org/10.1007/s00113-003-0590-0",
    "Lobenhoffer P, Gerich T, Bertram T, Lattermann C, Pohlemann T, Tscherne H. Particular surgical exposures for the treatment of tibial head fractures with posteromedial fragments. Unfallchirurg. 1997;100(12):957–967. https://doi.org/10.1007/s001130050213",
]


SURGICAL_TECHNIQUE_FIGURES = [
    "Figure 1.  Operative photograph showing the curved posteromedial skin incision behind the medial collateral ligament with the saphenous neurovascular bundle protected superiorly.",
    "Figure 2.  Operative photograph showing the exposed posteromedial fragment after capsulotomy, with the medial gastrocnemius retracted posteriorly.",
    "Figure 3.  Postoperative anteroposterior and lateral radiographs demonstrating anatomic reduction of the posteromedial fragment and stable buttress plate fixation.",
    "Figure 4.  Patient positioning diagram showing the supine position with hip bump and the location of the posteromedial incision relative to surface anatomy.",
]


def images_in_orthopedics_example_sections() -> List[dict]:
    return [
        main_heading("Keywords"),
        para(
            "calcaneal stress fracture; running injury; magnetic resonance "
            "imaging"
        ),
        main_heading("Clinical description", page_break_before=True),
        para_runs(
            "A 32-year-old recreational marathon runner presented with a "
            "six-week history of bilateral heel pain that had worsened "
            "progressively over the prior four weeks of training and was "
            "poorly relieved by rest. He had no preceding trauma and no "
            "prior history of stress fracture. His weekly running volume "
            "had increased from 35 to 65 miles over the eight weeks "
            "preceding presentation in preparation for an upcoming "
            "marathon. Physical examination demonstrated point tenderness "
            "over the posterior calcaneal tuberosities bilaterally without "
            "swelling or erythema. Plain radiographs were unremarkable. "
            "Magnetic resonance imaging demonstrated bilateral linear low-"
            "signal lesions in the posterior calcanei with surrounding "
            "marrow edema on T2-weighted sequences, consistent with "
            "bilateral calcaneal stress fractures (Figure 1)",
            cite(1),
            ".",
        ),
        para_runs(
            "The patient was placed in a controlled ankle-motion boot with "
            "progression to weight-bearing as tolerated and instructed to "
            "discontinue running for eight weeks. Vitamin D, calcium, and "
            "serum 25-hydroxyvitamin D levels were within normal limits. "
            "At eight-week follow-up, his pain had resolved and a graded "
            "return-to-run program was initiated",
            cite(2),
            ". He had no recurrence at six months.",
        ),
        main_heading("Teaching point"),
        para(
            "Calcaneal stress fractures should be considered in runners "
            "with recent training-volume escalation and posterior heel "
            "pain unresponsive to rest, even when plain radiographs are "
            "unremarkable. Magnetic resonance imaging is the imaging "
            "modality of choice for confirming the diagnosis at the pre-"
            "fracture-line stage."
        ),
    ]


IMAGES_REFERENCES = [
    "Sormaala MJ, Niva MH, Kiuru MJ, Mattila VM, Pihlajamäki HK. Stress injuries of the calcaneus detected with magnetic resonance imaging in military recruits. J Bone Joint Surg Am. 2006;88(10):2237–2242. https://doi.org/10.2106/JBJS.E.01190",
    "Kahanov L, Eberman LE, Games KE, Wasik M. Diagnosis, treatment, and rehabilitation of stress fractures in the lower extremity in runners. Open Access J Sports Med. 2015;6:87–95. https://doi.org/10.2147/OAJSM.S39512",
]


IMAGES_FIGURES = [
    "Figure 1.  Sagittal short-tau inversion recovery (STIR) magnetic resonance image of the right calcaneus demonstrating a linear low-signal stress fracture line in the posterior calcaneal tuberosity with surrounding marrow edema.",
]


def letter_to_editor_example_sections() -> List[dict]:
    return [
        main_heading("Body", page_break_before=True),
        para(
            "We read with interest the recent case series by Patel and "
            "colleagues on staged bilateral core decompression with bone "
            "marrow aspirate concentrate for COVID-era corticosteroid-"
            "induced femoral head osteonecrosis (OSCRSJ. 2026;1(2):e0007. "
            "https://doi.org/10.99999/oscrsj.2026.0007). The authors "
            "report excellent radiographic and functional outcomes at "
            "minimum 18-month follow-up in a population that had received "
            "cumulative corticosteroid exposure substantially below the "
            "historical 2,000 mg prednisone-equivalent threshold. We "
            "commend the authors on a careful series and offer two "
            "observations."
        ),
        para_runs(
            "First, the authors note that the single conversion to total "
            "hip arthroplasty in their cohort occurred in the patient with "
            "the largest preoperative lesion volume (47% of femoral head). "
            "This finding is consistent with the broader literature on "
            "lesion size as a predictor of progression after core "
            "decompression",
            cite(1, 2),
            ". We would suggest that future series stratify outcomes by "
            "preoperative lesion volume in 10% or 25% volumetric "
            "increments rather than reporting a continuous variable, "
            "which would aid surgical decision-making at the bedside.",
        ),
        para_runs(
            "Second, the authors do not address whether biologic "
            "augmentation with bone marrow aspirate concentrate is "
            "responsible for the reported outcomes beyond what would be "
            "expected with core decompression alone. The literature on "
            "this question remains inconclusive. A recent meta-analysis by "
            "Hernigou and colleagues reported a 12-percentage-point "
            "reduction in conversion to arthroplasty with BMAC "
            "augmentation versus core decompression alone, but "
            "heterogeneity in injection technique and concentration limits "
            "the strength of that finding",
            cite(3),
            ". A multicenter randomized comparison of core decompression "
            "alone versus core decompression with BMAC in this specific "
            "COVID-era population would be the natural next step and would "
            "help clarify whether BMAC is an essential adjunct or a useful "
            "but optional one.",
        ),
        para(
            "We thank the authors for their contribution and look forward "
            "to further reports from this cohort at extended follow-up."
        ),
    ]


LETTER_REFERENCES = [
    "Mont MA, Cherian JJ, Sierra RJ, Jones LC, Lieberman JR. Nontraumatic osteonecrosis of the femoral head: where do we stand today? J Bone Joint Surg Am. 2015;97(19):1604–1627. https://doi.org/10.2106/JBJS.O.00071",
    "Steinberg ME, Hayken GD, Steinberg DR. A quantitative system for staging avascular necrosis. J Bone Joint Surg Br. 1995;77(1):34–41.",
    "Hernigou P, Beaujean F. Treatment of osteonecrosis with autologous bone marrow grafting. Clin Orthop Relat Res. 2002;405:14–23. https://doi.org/10.1097/00003086-200212000-00003",
]


LETTER_FIGURES: List[str] = []  # no figures in this example


def review_article_example_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        sub_heading("Background"),
        para(
            "Bone marrow aspirate concentrate (BMAC) has been proposed as "
            "a biologic augmentation strategy for joint preservation in "
            "pre-collapse femoral head osteonecrosis. The strength of "
            "evidence supporting routine BMAC use is contested."
        ),
        sub_heading("Methods"),
        para(
            "We performed a systematic review of MEDLINE, Embase, and the "
            "Cochrane Library through December 2025 for studies comparing "
            "core decompression alone with core decompression plus BMAC in "
            "pre-collapse (Ficat I or II) femoral head osteonecrosis. "
            "Outcomes of interest were conversion to total hip "
            "arthroplasty and Harris Hip Score at minimum 24 months."
        ),
        sub_heading("Results"),
        para(
            "Twelve studies (1,247 hips) met inclusion criteria. Pooled "
            "conversion to arthroplasty was 18% with core decompression "
            "alone versus 11% with BMAC augmentation (relative risk 0.61; "
            "95% confidence interval 0.48–0.78). Mean improvement in "
            "Harris Hip Score was modestly greater in the BMAC arm (+8 "
            "points; p = 0.02). Heterogeneity in injection technique, "
            "BMAC concentration, and lesion staging limited the strength "
            "of effect-size estimates."
        ),
        sub_heading("Conclusion"),
        para(
            "BMAC augmentation may reduce conversion to arthroplasty in "
            "pre-collapse femoral head osteonecrosis, but heterogeneity in "
            "technique limits clinical generalizability. Multicenter "
            "prospective trials with standardized BMAC preparation are "
            "needed."
        ),
        main_heading("Keywords"),
        para(
            "avascular necrosis; femoral head; core decompression; bone "
            "marrow aspirate concentrate; systematic review"
        ),
        main_heading("Introduction", page_break_before=True),
        para_runs(
            "Femoral head osteonecrosis is a debilitating condition that, "
            "untreated, progresses to subchondral collapse and end-stage "
            "osteoarthritis in more than 80% of cases",
            cite(1),
            ". Joint-preserving interventions in the pre-collapse stage "
            "have the potential to delay or obviate total hip arthroplasty "
            "in young patients, who otherwise face the prospect of one or "
            "more revision arthroplasties over their remaining lifespan",
            cite(2),
            ".",
        ),
        para_runs(
            "Core decompression remains the most extensively studied "
            "joint-preserving intervention. Reported success rates in "
            "Ficat stage I and II disease range from 60% to 80% across "
            "single- and multi-center series",
            cite(3, 4),
            ". Biologic augmentation strategies — including autologous "
            "bone marrow aspirate concentrate (BMAC), demineralized bone "
            "matrix, and platelet-rich plasma — have emerged over the "
            "past two decades as proposed adjuncts to core decompression. "
            "The strength of evidence supporting routine biologic "
            "augmentation is contested, with proponents citing improved "
            "early outcomes and skeptics citing methodological "
            "heterogeneity and small effect sizes.",
        ),
        para(
            "We performed a systematic review of the comparative "
            "literature evaluating core decompression alone versus core "
            "decompression with BMAC augmentation in pre-collapse femoral "
            "head osteonecrosis to clarify the strength and consistency "
            "of any reported clinical benefit."
        ),
        main_heading("Methods"),
        para(
            "We conducted a systematic search of MEDLINE (via PubMed), "
            "Embase, and the Cochrane Central Register of Controlled "
            "Trials through December 31, 2025. The search strategy "
            "combined terms for femoral head osteonecrosis, core "
            "decompression, and bone marrow aspirate concentrate. "
            "Inclusion criteria were: (1) randomized controlled trial or "
            "comparative cohort study; (2) pre-collapse (Ficat I or II) "
            "disease; (3) explicit comparison of core decompression alone "
            "versus core decompression with BMAC; (4) minimum 24-month "
            "follow-up; (5) English-language publication."
        ),
        para(
            "Two reviewers independently screened titles and abstracts; "
            "disagreements were resolved by consensus. Quality of included "
            "studies was assessed using the Cochrane Risk of Bias 2 tool "
            "for randomized trials and the Newcastle-Ottawa Scale for "
            "cohort studies. The primary outcome was conversion to total "
            "hip arthroplasty by final follow-up. Secondary outcomes "
            "included Harris Hip Score change from baseline. Random-"
            "effects meta-analysis was used to pool effect sizes due to "
            "anticipated heterogeneity in injection technique and BMAC "
            "preparation."
        ),
        main_heading("Results"),
        para(
            "The initial search yielded 384 records. After deduplication "
            "and title/abstract screening, 47 full-text articles were "
            "assessed; 12 studies (1,247 hips in 1,084 patients) met "
            "inclusion criteria. Patient and study characteristics are "
            "summarized in Table 1."
        ),
        para(
            "Pooled conversion to total hip arthroplasty was 18.2% (95% "
            "confidence interval 14.7–22.1%) with core decompression alone "
            "versus 11.0% (95% confidence interval 8.4–14.1%) with BMAC "
            "augmentation. The pooled relative risk for conversion was "
            "0.61 (95% confidence interval 0.48–0.78; I² = 34%). Mean "
            "improvement in Harris Hip Score from baseline to final "
            "follow-up was +18 points with core decompression alone and "
            "+26 points with BMAC augmentation (mean difference +8 "
            "points; 95% confidence interval +1 to +15; p = 0.02)."
        ),
        para(
            "Heterogeneity in BMAC concentration ranged from 4× to 24× "
            "nucleated cell concentration relative to peripheral marrow. "
            "Injection technique varied from single-bolus injection "
            "through the trephine track to multiple smaller injections at "
            "multiple decompression channels. Subgroup analysis by lesion "
            "stage demonstrated a more pronounced benefit of BMAC in "
            "Ficat stage II disease (relative risk 0.55) than in Ficat "
            "stage I disease (relative risk 0.74)."
        ),
        main_heading("Discussion"),
        para_runs(
            "This systematic review identifies a modest but consistent "
            "reduction in conversion to total hip arthroplasty with BMAC "
            "augmentation of core decompression in pre-collapse femoral "
            "head osteonecrosis. The effect size is consistent with prior "
            "meta-analyses",
            cite(5, 6),
            " and is more pronounced in Ficat stage II disease, where the "
            "underlying probability of progression with core decompression "
            "alone is higher.",
        ),
        para(
            "Three methodological limitations temper the strength of any "
            "single-study or pooled estimate. First, BMAC concentration "
            "and injection technique varied substantially across the "
            "included studies. Second, lesion staging in the included "
            "literature relied on Ficat classification, which is known to "
            "underestimate lesion volume relative to MRI-based volumetric "
            "assessment. Third, the included literature is dominated by "
            "single-center retrospective cohorts; only three of the 12 "
            "included studies were prospective randomized trials."
        ),
        para(
            "A multicenter randomized comparison of core decompression "
            "alone versus core decompression with BMAC in a population "
            "stratified by preoperative lesion volume — using a "
            "standardized BMAC preparation protocol — would address all "
            "three of these limitations and would provide the strongest "
            "evidence base for routine biologic augmentation in this "
            "population."
        ),
        main_heading("Conclusion"),
        para(
            "BMAC augmentation of core decompression may reduce conversion "
            "to total hip arthroplasty in pre-collapse femoral head "
            "osteonecrosis, with a more pronounced effect in Ficat stage "
            "II disease. Heterogeneity in BMAC preparation and injection "
            "technique limits the generalizability of pooled estimates. "
            "Standardized multicenter trials are needed."
        ),
    ]


REVIEW_REFERENCES = [
    "Mont MA, Cherian JJ, Sierra RJ, Jones LC, Lieberman JR. Nontraumatic osteonecrosis of the femoral head: where do we stand today? J Bone Joint Surg Am. 2015;97(19):1604–1627. https://doi.org/10.2106/JBJS.O.00071",
    "Lieberman JR, Berry DJ, Mont MA, et al. Osteonecrosis of the hip: management in the 21st century. Instr Course Lect. 2003;52:337–355.",
    "Marker DR, Seyler TM, Ulrich SD, Srivastava S, Mont MA. Do modern techniques improve core decompression outcomes for hip osteonecrosis? Clin Orthop Relat Res. 2008;466(5):1093–1103. https://doi.org/10.1007/s11999-008-0184-9",
    "Pierce TP, Jauregui JJ, Cherian JJ, Elmallah RK, Mont MA. Outcomes of total hip arthroplasty in patients with osteonecrosis of the femoral head. Clin Orthop Relat Res. 2015;473(8):2487–2493. https://doi.org/10.1007/s11999-015-4252-7",
    "Hernigou P, Beaujean F. Treatment of osteonecrosis with autologous bone marrow grafting. Clin Orthop Relat Res. 2002;405:14–23. https://doi.org/10.1097/00003086-200212000-00003",
    "Andriolo L, Merli G, Tobar C, Altamura SA, Kon E, Filardo G. Cell therapy for osteonecrosis of the femoral head: a systematic review. Stem Cells Int. 2018;2018:4548349. https://doi.org/10.1155/2018/4548349",
]


REVIEW_FIGURES = [
    "Figure 1.  PRISMA flow diagram of study identification, screening, eligibility, and inclusion through the systematic review process.",
    "Figure 2.  Forest plot of pooled relative risk of conversion to total hip arthroplasty across the 12 included studies.",
]


# ---------------------------------------------------------------------------
# Example specs
# ---------------------------------------------------------------------------


EXAMPLES = [
    {
        "filename": "oscrsj-example-case-report.docx",
        "running_title": "Iatrogenic median nerve injury — supracondylar pinning",
        "sections": case_report_example_sections,
        "references": CASE_REPORT_REFERENCES,
        "figures": CASE_REPORT_FIGURES,
    },
    {
        "filename": "oscrsj-example-case-series.docx",
        "running_title": "Core decompression with BMAC for COVID-era AVN",
        "sections": case_series_example_sections,
        "references": CASE_SERIES_REFERENCES,
        "figures": CASE_SERIES_FIGURES,
    },
    {
        "filename": "oscrsj-example-surgical-technique.docx",
        "running_title": "Modified posteromedial approach for tibial plateau",
        "sections": surgical_technique_example_sections,
        "references": SURGICAL_TECHNIQUE_REFERENCES,
        "figures": SURGICAL_TECHNIQUE_FIGURES,
    },
    {
        "filename": "oscrsj-example-images-in-orthopedics.docx",
        "running_title": "Bilateral calcaneal stress fractures in a runner",
        "sections": images_in_orthopedics_example_sections,
        "references": IMAGES_REFERENCES,
        "figures": IMAGES_FIGURES,
    },
    {
        "filename": "oscrsj-example-letter-to-editor.docx",
        "running_title": "Letter — re: Patel et al. on BMAC for COVID-era AVN",
        "sections": letter_to_editor_example_sections,
        "references": LETTER_REFERENCES,
        "figures": LETTER_FIGURES,
    },
    {
        "filename": "oscrsj-example-review-article.docx",
        "running_title": "BMAC vs core decompression alone — systematic review",
        "sections": review_article_example_sections,
        "references": REVIEW_REFERENCES,
        "figures": REVIEW_FIGURES,
    },
]


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------


def render_example(spec: dict) -> Path:
    doc = Document()
    configure_document_defaults(doc)
    section = doc.sections[0]
    configure_section(section, running_title=spec["running_title"])

    # ----- main manuscript body -----
    sections = spec["sections"]()
    is_first_block = True
    for entry in sections:
        kind = entry["kind"]
        if kind == "h1":
            if not is_first_block:
                add_blank_paragraph(doc)
            add_heading_paragraph(
                doc,
                entry["text"],
                centered=False,
                page_break_before=entry.get("page_break_before", False),
            )
            is_first_block = False
        elif kind == "h2":
            add_blank_paragraph(doc)
            add_heading_paragraph(doc, entry["text"], centered=False)
            is_first_block = False
        elif kind == "para":
            add_paragraph(doc, entry["text"])
            is_first_block = False
        elif kind == "para_runs":
            add_paragraph_with_citations(doc, entry["segments"])
            is_first_block = False
        else:
            raise ValueError(f"Unknown section kind: {kind!r}")

    # ----- References (new page, centred sentence-case heading, no blank line below) -----
    add_heading_paragraph(doc, "References", centered=True, page_break_before=True)
    for n, ref_text in enumerate(spec["references"], start=1):
        ref_para = add_paragraph(
            doc,
            f"{n}.  {ref_text}",
            indent_hanging_inches=0.25,
        )
        add_bookmark(ref_para, f"REF_{n}")

    # ----- Figure legends (new page, centred sentence-case heading, no blank line below) -----
    if spec["figures"]:
        add_heading_paragraph(
            doc, "Figure legends", centered=True, page_break_before=True
        )
        for legend in spec["figures"]:
            add_paragraph(doc, legend)

    out_path = DOWNLOADS_DIR / spec["filename"]
    doc.save(str(out_path))
    return out_path


def main() -> None:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Writing worked examples to: {DOWNLOADS_DIR}\n")
    for spec in EXAMPLES:
        path = render_example(spec)
        size_kb = path.stat().st_size / 1024
        print(f"  ✓ {spec['filename']:<48} {size_kb:6.1f} KB")
    print("\nDone.")


if __name__ == "__main__":
    main()
