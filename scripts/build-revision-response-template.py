#!/usr/bin/env python3
"""
Build the OSCRSJ Revision Response Template .docx at
public/downloads/oscrsj-revision-response-template.docx.

Run from the repo root: `python3 scripts/build-revision-response-template.py`

Requires `python-docx` (`pip install python-docx`).

Authors download this template when they receive a Minor or Major Revisions
decision. They use it to assemble the point-by-point response document that is
required (alongside the tracked-changes manuscript with red-text /
yellow-highlighted edits) when re-submitting through the revision wizard.

Structure:
  Cover instructions — what the document is, what's required, how to format
                       changes in the manuscript.
  Section 1 — Manuscript metadata block (manuscript ID / title / corresponding
              author / date / decision).
  Section 2 — How to use this template (numbered steps).
  Section 3 — WORKED EXAMPLE (3 sample reviewer comments fully populated).
  Section 4 — BLANK RESPONSE TABLE (3 starter rows authors expand as needed).
  Section 5 — Submission checklist + footer.

History:
  v1.0 (2026-04-26, Session 36) — initial release.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "public" / "downloads" / "oscrsj-revision-response-template.docx"

BROWN_DARK = RGBColor(0x3D, 0x2A, 0x18)
INK = RGBColor(0x12, 0x0D, 0x08)
TAN = RGBColor(0x99, 0x7E, 0x67)
PEACH_DARK = RGBColor(0xF0, 0xC4, 0x9A)
RED = RGBColor(0xC0, 0x39, 0x2B)
BORDER = "997E67"
HEADER_FILL = "F8F4ED"  # cream-alt
ROW_ALT_FILL = "FDFBF8"  # cream
EXAMPLE_FILL = "F7F6F4"


def set_cell_border(cell, color: str = BORDER, sz: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        tc_borders.append(e)
    tc_pr.append(tc_borders)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(par, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6, line=1.15)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = BROWN_DARK
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def add_para(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    bold: bool = False,
    size: int = 10,
    after: int = 6,
    color: RGBColor | None = None,
) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else INK


def add_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1.2)
    run = p.add_run(label)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = TAN


def add_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    color: RGBColor | None = None,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Replace a cell's content with a single styled paragraph."""
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.25)
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else INK


def append_cell_para(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    color: RGBColor | None = None,
) -> None:
    """Append an additional paragraph inside an existing cell."""
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1.25)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color is not None else INK


def build_response_table(
    doc: Document,
    rows: list[tuple[str, str, str, str, str]],
    *,
    fill_alt: bool = False,
) -> None:
    """
    Build a 4-column response table.
    Columns: # | Reviewer Comment | Author Response | Change Made & Line Numbers
    `rows` is a list of (number, reviewer_comment, author_response, change_made, line_ref).
    `change_made` and `line_ref` are concatenated into the 4th column.
    """
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.autofit = False
    widths = [Cm(1.0), Cm(5.5), Cm(5.5), Cm(4.5)]
    for col, w in zip(table.columns, widths):
        col.width = w

    # Header row
    headers = ("#", "Reviewer Comment", "Author Response", "Change Made & Line Numbers")
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell)
        set_cell_shading(cell, HEADER_FILL)
        add_cell_text(cell, label, bold=True, size=10, color=BROWN_DARK)

    # Data rows
    for r, (num, comment, response, change, line_ref) in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, cell in enumerate(cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell)
            if fill_alt and r % 2 == 0:
                set_cell_shading(cell, ROW_ALT_FILL)

        # Column 1 — #
        add_cell_text(cells[0], num, bold=True, size=11, color=BROWN_DARK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # Column 2 — Reviewer Comment
        if comment:
            add_cell_text(cells[1], comment, size=10)
        else:
            add_cell_text(cells[1], "[Paste reviewer comment verbatim]",
                          italic=True, size=10, color=TAN)

        # Column 3 — Author Response
        if response:
            add_cell_text(cells[2], response, size=10)
        else:
            add_cell_text(cells[2],
                          "[Your point-by-point response. Acknowledge the comment, "
                          "explain how you addressed it, or — if you disagree — explain why.]",
                          italic=True, size=10, color=TAN)

        # Column 4 — Change Made + Line Numbers
        if change:
            add_cell_text(cells[3], change, size=10)
            if line_ref:
                append_cell_para(cells[3], line_ref, italic=True, size=9, color=TAN)
        else:
            add_cell_text(cells[3],
                          "[Describe the specific change made to the manuscript.]",
                          italic=True, size=10, color=TAN)
            append_cell_para(cells[3],
                             "[Line numbers, e.g., L142–148]",
                             italic=True, size=9, color=TAN)


def build() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = INK

    # ===== Header =====
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSCRSJ")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BROWN_DARK

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Orthopedic Surgery Case Reports & Series Journal")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = BROWN_DARK

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REVISION RESPONSE TEMPLATE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TAN

    add_para(
        doc,
        "Use this document to respond to reviewer comments after receiving a "
        "Minor or Major Revisions decision. Submit the completed file as your "
        "Response to Reviewers in Step 2 of the revision wizard, alongside your "
        "tracked-changes manuscript.",
        italic=True,
        size=10,
        after=14,
    )

    # ===== Section 1 — Manuscript metadata =====
    add_heading(doc, "1. Manuscript Information", level=1)

    fields = [
        ("Manuscript ID", "[OSCRSJ-2026-####]"),
        ("Manuscript Title", "[Title as it appears on your submission]"),
        ("Corresponding Author", "[Your full name]"),
        ("Decision Received", "[Minor Revisions / Major Revisions]"),
        ("Date of Resubmission", "[YYYY-MM-DD]"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11.0)
    for i, (label, placeholder) in enumerate(fields):
        cells = table.rows[i].cells
        cells[0].width = Cm(5.5)
        cells[1].width = Cm(11.0)
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        for cell in cells[0].paragraphs:
            cell.paragraph_format.space_after = Pt(2)
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BROWN_DARK
        run = cells[1].paragraphs[0].add_run(placeholder)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = TAN

    # spacer
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)

    # ===== Section 2 — How to use this template =====
    add_heading(doc, "2. How to Use This Template", level=1)
    add_para(
        doc,
        "Every revision submitted to OSCRSJ requires TWO documents in addition "
        "to the clean revised manuscript:",
        bold=True,
    )

    bullets = [
        (
            "(a) Response to Reviewers (this document).",
            "A point-by-point response. For every reviewer comment, paste the "
            "comment verbatim, write your response, describe the specific change "
            "you made to the manuscript, and cite the line numbers in the "
            "revised manuscript where the change appears. Use Section 4 below.",
        ),
        (
            "(b) Tracked-Changes Manuscript.",
            "The revised manuscript with every change visually marked. New or "
            "changed text must be in RED font and HIGHLIGHTED YELLOW. "
            "Deletions can be shown via Word's track-changes feature or by "
            "leaving the deleted text struck-through in red. The editor must "
            "be able to see — at a glance — every change you made.",
        ),
    ]
    for label, body in bullets:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.3)
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BROWN_DARK
        run = p.add_run(body)
        run.font.size = Pt(10)
        run.font.color.rgb = INK

    add_para(
        doc,
        "Steps:",
        bold=True,
        after=4,
    )
    steps = [
        "Read the editor's decision letter and every reviewer's comments in full before editing the manuscript.",
        "Open your manuscript in Word. Make every revision in RED font and HIGHLIGHT each change YELLOW. Save this as your tracked-changes file.",
        "Save a second copy with all formatting cleaned (no red, no highlighting). This is your clean revised manuscript and revised blinded manuscript.",
        "Fill out the response table in Section 4 below. One row per reviewer comment. Quote the comment verbatim — do not paraphrase.",
        "For each comment, write a clear response, describe the specific change, and cite the line numbers in your REVISED manuscript (not the original).",
        "If you disagree with a comment, say so respectfully and explain your reasoning. Disagreement is acceptable; ignoring a comment is not.",
        "Upload all three files in Step 2 of the revision wizard: revised manuscript (clean), revised blinded manuscript (clean), tracked-changes manuscript, and this Response to Reviewers document.",
    ]
    for n, body in enumerate(steps, start=1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=4, line=1.3)
        run = p.add_run(f"{n}.  ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BROWN_DARK
        run = p.add_run(body)
        run.font.size = Pt(10)
        run.font.color.rgb = INK

    # ===== Section 3 — Worked example =====
    add_heading(doc, "3. Worked Example", level=1)
    add_para(
        doc,
        "The example below shows three reviewer comments fully populated. "
        "Your response should follow the same structure.",
        italic=True,
    )

    # Editor section
    add_heading(doc, "Editor's Comments", level=2)
    example_editor_rows = [
        (
            "E1",
            "Please ensure the manuscript adheres to the CARE checklist. The "
            "timeline of follow-up needs clarification, and the discussion would "
            "benefit from comparison with at least two recent published reports.",
            "Thank you. We have completed the CARE checklist (now uploaded as a "
            "supplementary file). The follow-up timeline has been clarified in the "
            "Case Presentation, and we have added comparison with Smith et al. "
            "(2024) and Garcia et al. (2025) in the Discussion.",
            "Updated CARE checklist; revised follow-up paragraph in Case "
            "Presentation; added two-paragraph comparison in Discussion.",
            "L88–94 (follow-up); L210–238 (comparison)",
        ),
    ]
    build_response_table(doc, example_editor_rows)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)

    # Reviewer 1
    add_heading(doc, "Reviewer 1", level=2)
    example_r1_rows = [
        (
            "R1.1",
            "The introduction does not adequately establish why this case is "
            "novel. The cited references on talar fractures are from 2009 and 2012; "
            "more recent literature should be incorporated.",
            "Agreed. We have rewritten the second paragraph of the Introduction "
            "to explicitly state the novelty (combination of three-column "
            "involvement and a non-anatomic fixation strategy not previously "
            "reported). We replaced the dated references with four recent "
            "publications (2022–2025).",
            "Rewrote Introduction paragraph 2; updated references 4–7.",
            "L42–58 (introduction); references 4, 5, 6, 7",
        ),
        (
            "R1.2",
            "Figure 2 is low resolution and the arrow markers are difficult to "
            "see. Please re-export at higher DPI.",
            "We re-exported Figure 2 at 600 DPI from the original DICOM source "
            "and replaced the arrow markers with a higher-contrast color (peach "
            "on dark background) for clearer visibility.",
            "Replaced Figure_2.tiff with a 600 DPI re-export; new annotations.",
            "Figure 2 (file replaced)",
        ),
        (
            "R1.3",
            "The Discussion claims long-term outcome data but follow-up is only "
            "9 months. Please soften this claim.",
            "Agreed. We have removed the phrase \"long-term outcome\" and replaced "
            "it with \"short-term functional outcome at 9-month follow-up\" "
            "throughout the Discussion. We have added a Limitations subsection "
            "explicitly noting the short follow-up duration.",
            "Replaced 4 instances of \"long-term outcome\" with the corrected "
            "phrasing; added Limitations subsection.",
            "L245, L258, L271, L286 (replacements); L292–301 (limitations)",
        ),
    ]
    build_response_table(doc, example_r1_rows, fill_alt=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)

    # Reviewer 2
    add_heading(doc, "Reviewer 2", level=2)
    example_r2_rows = [
        (
            "R2.1",
            "The conclusion overreaches given a single case. Recommend reframing "
            "as hypothesis-generating rather than evidence for a new technique.",
            "We respectfully agree in part. We have reframed the Conclusion to "
            "describe our finding as \"hypothesis-generating\" and have removed the "
            "phrase \"new technique.\" However, we have retained the description "
            "of the surgical approach taken, since this is the central "
            "educational value of the case report.",
            "Rewrote the Conclusion (3 sentences); removed \"new technique\" "
            "throughout; preserved the technique description.",
            "L312–319 (conclusion)",
        ),
    ]
    build_response_table(doc, example_r2_rows)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=10)
    run = p.add_run(
        "End of worked example. Begin your own response below."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = TAN

    # ===== Section 4 — Blank response table =====
    add_heading(doc, "4. Your Response", level=1)
    add_para(
        doc,
        "Add a separate sub-section for the editor and each reviewer (R1, R2, "
        "R3, etc.). Number every comment within each reviewer's section "
        "(E1, E2, R1.1, R1.2, R2.1, etc.). Add as many rows as you need by "
        "right-clicking inside the table and choosing Insert → Rows Below.",
        italic=True,
    )

    add_heading(doc, "Editor's Comments", level=2)
    blank_editor_rows = [
        ("E1", "", "", "", ""),
    ]
    build_response_table(doc, blank_editor_rows)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)

    add_heading(doc, "Reviewer 1", level=2)
    blank_r1_rows = [
        ("R1.1", "", "", "", ""),
        ("R1.2", "", "", "", ""),
        ("R1.3", "", "", "", ""),
    ]
    build_response_table(doc, blank_r1_rows, fill_alt=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)

    add_heading(doc, "Reviewer 2", level=2)
    blank_r2_rows = [
        ("R2.1", "", "", "", ""),
        ("R2.2", "", "", "", ""),
        ("R2.3", "", "", "", ""),
    ]
    build_response_table(doc, blank_r2_rows, fill_alt=True)

    add_para(
        doc,
        "Add a Reviewer 3 / Reviewer 4 section below by copying the heading + "
        "table format above. Every reviewer comment in your decision letter must "
        "be answered.",
        italic=True,
        after=10,
    )

    # ===== Section 5 — Submission checklist =====
    add_heading(doc, "5. Pre-Submission Checklist", level=1)
    add_para(
        doc,
        "Before re-submitting through the OSCRSJ revision wizard, confirm that "
        "all the following are true.",
        italic=True,
    )

    checklist = [
        "Every reviewer comment has been answered. No comments are skipped.",
        "Each comment is quoted verbatim — not paraphrased.",
        "Each response cites specific line numbers in the REVISED manuscript.",
        "Where I disagreed with a comment, I explained my reasoning respectfully.",
        "The tracked-changes manuscript has every change in RED font and HIGHLIGHTED YELLOW.",
        "The clean revised manuscript has all formatting (red, highlighting, struck-through text) removed.",
        "The clean revised BLINDED manuscript has identifying information removed AND all change-tracking removed.",
        "I have NOT used AI to write this response document or the manuscript revisions.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=3, line=1.3)
        run = p.add_run("☐  ")
        run.font.size = Pt(11)
        run.font.color.rgb = BROWN_DARK
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = INK

    # ===== Footer =====
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OSCRSJ Editorial Office  ·  editorial@oscrsj.com  ·  www.oscrsj.com")
    run.font.size = Pt(9)
    run.font.color.rgb = TAN

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Revision Response Template v1.0  ·  Issued April 2026")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = TAN

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(REPO_ROOT)} ({OUTPUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
