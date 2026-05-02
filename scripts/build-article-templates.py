#!/usr/bin/env python3
"""
Build all 6 OSCRSJ article-type manuscript templates as .docx files at
public/downloads/oscrsj-template-{type}.docx.

Run from the repo root: `python3 scripts/build-article-templates.py`

Requires `python-docx` (`pip install python-docx`).

This is the unified, paste-runnable regenerator for the article-type manuscript
templates. The script defines a TEMPLATES config dict that lists the section
structure of each template; the same renderer applies to every template so
every template gets the same headings, header, page-break, and citation
treatment.

Locked formatting decisions (Kanwar directive 2026-05-01):

  1. All headings render in sentence case (only the first letter capitalised).
  2. Introduction (the main-manuscript Introduction) starts on a new page.
  3. Every heading has an empty line ABOVE it, never below.
  4. No italics anywhere in the headings.
  5. The document carries a header with the author's short running title
     anchored to the top-right corner of every page.
  6. In-text citations render as superscript Word hyperlinks pointing to the
     matching reference at the end of the manuscript. The template demonstrates
     this with a worked example in the Introduction and seeds bookmarks at
     References [1] / [2] / [3] so the example links are live.
  7. References sits on its own page. The "References" heading is centred,
     sentence case, and there is no empty paragraph between the heading and the
     first reference.
  8. Figure legends mirrors References — its own page, centred sentence-case
     heading, no empty paragraph between the heading and the first legend.

History:
  v1.2 (2026-05-01) — initial unified release covering all 6 article types.
                       Supersedes the 2026-04-25 single-template build path
                       captured in the v1.1 template ship (Session 19).
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Emu, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths / fonts / colour constants
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
DOWNLOADS_DIR = REPO_ROOT / "public" / "downloads"

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
LINE_SPACING = 2.0  # double-spaced
BLACK = RGBColor(0x00, 0x00, 0x00)
HYPERLINK_BLUE = RGBColor(0x05, 0x63, 0xC1)


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------


def _set_default_run_font(run) -> None:
    run.font.name = FONT_NAME
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    run.font.size = Pt(FONT_SIZE_PT)
    run.font.color.rgb = BLACK


def _apply_paragraph_defaults(par, *, double_spaced: bool = True) -> None:
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if double_spaced:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_blank_paragraph(doc, *, in_section=None) -> None:
    """Add an empty paragraph (used to put a blank line ABOVE a heading)."""
    target = in_section if in_section is not None else doc
    p = target.add_paragraph()
    _apply_paragraph_defaults(p)


def add_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: Optional[int] = None,
    indent_hanging_inches: Optional[float] = None,
    double_spaced: bool = True,
) -> object:
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, double_spaced=double_spaced)
    if align is not None:
        p.alignment = align
    if indent_hanging_inches is not None:
        pf = p.paragraph_format
        pf.left_indent = Pt(36)
        pf.first_line_indent = Pt(-36)
    run = p.add_run(text)
    _set_default_run_font(run)
    run.bold = bold
    run.italic = italic
    return p


def add_heading_paragraph(
    doc,
    text: str,
    *,
    centered: bool = False,
    page_break_before: bool = False,
) -> object:
    """Bold, sentence-case, no italics. Optionally centered, optionally on a new page."""
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if page_break_before:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    _set_default_run_font(run)
    run.bold = True
    run.italic = False
    return p


def add_bookmark(paragraph, name: str) -> None:
    """Insert a Word bookmark wrapping the START of the paragraph's first run.

    Bookmark IDs are kept globally unique by walking the document's existing
    bookmark IDs at call time.
    """
    doc = paragraph._parent
    body = doc._element
    existing_ids = [
        int(b.get(qn("w:id")))
        for b in body.iter(qn("w:bookmarkStart"))
        if b.get(qn("w:id")) is not None
    ]
    next_id = (max(existing_ids) + 1) if existing_ids else 0

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(next_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(next_id))

    p = paragraph._p
    # Insert <w:bookmarkStart> at the top of the paragraph (before any runs)
    # and <w:bookmarkEnd> immediately after, so the bookmark anchors at the
    # start of the reference line. Hyperlink jumps land at this point.
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(end)
        pPr.addnext(start)
    else:
        p.insert(0, end)
        p.insert(0, start)


def add_internal_hyperlink_run(
    paragraph,
    text: str,
    bookmark_name: str,
    *,
    superscript: bool = True,
) -> None:
    """Append a run linked to an internal bookmark; styled blue + superscript."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rPr.append(rfonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(FONT_SIZE_PT * 2))
    rPr.append(sz)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if superscript:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        rPr.append(vert)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_paragraph_with_inline_superscripts(
    doc,
    segments: List[Tuple[str, Optional[str]]],
    *,
    double_spaced: bool = True,
) -> object:
    """Create a paragraph by appending runs in order. Each segment is
    (text, bookmark_name_or_None). When bookmark_name is None the segment
    renders as plain text; when it's set, the segment renders as a superscript
    hyperlink anchored to that bookmark.
    """
    p = doc.add_paragraph()
    _apply_paragraph_defaults(p, double_spaced=double_spaced)
    for text, bookmark in segments:
        if bookmark is None:
            run = p.add_run(text)
            _set_default_run_font(run)
        else:
            add_internal_hyperlink_run(p, text, bookmark, superscript=True)
    return p


# ---------------------------------------------------------------------------
# Section-wide plumbing: header, line numbering, margins
# ---------------------------------------------------------------------------


def configure_section(section, *, running_title_placeholder: str) -> None:
    """Apply page margins, continuous line numbering, and the right-aligned
    running-title header to a section.
    """
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)

    sect_pr = section._sectPr
    # Continuous line numbering, restart on every page → use w:lnNumType
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:start"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    sect_pr.append(ln)

    # Right-aligned running title header on every page
    header = section.header
    # Clear any default empty paragraph that python-docx adds.
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(running_title_placeholder)
    _set_default_run_font(run)
    run.italic = True


def configure_document_defaults(doc) -> None:
    """Set the document-level default font + size so a casual user editing the
    template doesn't accidentally drop text in a different face."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)
    normal.font.color.rgb = BLACK
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)


# ---------------------------------------------------------------------------
# Template config — the section structure of every article type
# ---------------------------------------------------------------------------


# Each main-manuscript section is defined as (heading, [paragraphs]).
# A heading is rendered in sentence case with an empty paragraph above it.
# The Introduction heading is rendered with page_break_before so the body of
# the manuscript begins on a fresh page after the abstract + keywords.

# Subsection headings inside the structured abstract use the same rule
# (sentence case, empty line above, no italic).

INTRO_DEMO = (
    "[Replace this paragraph with your introduction content. Format every "
    "in-text citation as a superscript number that links to the matching "
    "entry in your Reference list. The next sentence shows the convention. "
    "Recent series have reported similar findings,"
)

# Marker used to splice the superscript hyperlinks into the demo intro.
# The renderer below reads each main-manuscript section and, for the very
# first section after the Keywords block (the Introduction), appends a
# 3-citation worked-example sentence with live superscript hyperlinks to
# the seeded References [1] / [2] / [3] bookmarks.

CITATION_HOWTO = (
    "[How to format in-text citations: select the citation number(s) in your "
    "manuscript, press Ctrl+Shift+= (Cmd+Shift+= on Mac) to apply Superscript, "
    "then Insert → Hyperlink → Place in This Document → REF_1 (or REF_2, "
    "REF_3, …) to link the superscript to the matching reference. Bookmarks "
    "REF_1, REF_2, and REF_3 are pre-seeded for you in the Reference list "
    "below; add REF_4, REF_5, … as you add references.]"
)


def main_paragraph(text: str) -> dict:
    return {"kind": "para", "text": text}


def main_heading(text: str, *, page_break_before: bool = False) -> dict:
    return {"kind": "h1", "text": text, "page_break_before": page_break_before}


def sub_heading(text: str) -> dict:
    return {"kind": "h2", "text": text}


# Section structure per template. Each list entry is either a heading dict or
# a paragraph dict. Headings are emitted in sentence case; paragraphs are
# rendered in Times New Roman 12pt double-spaced. The renderer inserts an
# empty paragraph BEFORE every heading (per requirement 3).


def case_report_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        main_paragraph("[Replace with your structured abstract — maximum 300 words.]"),
        sub_heading("Introduction"),
        main_paragraph("[Replace with your introduction content for this section of the abstract.]"),
        sub_heading("Case presentation"),
        main_paragraph("[Replace with your case presentation content for this section of the abstract.]"),
        sub_heading("Discussion"),
        main_paragraph("[Replace with your discussion content for this section of the abstract.]"),
        sub_heading("Conclusion"),
        main_paragraph("[Replace with your conclusion content for this section of the abstract.]"),
        main_heading("Keywords"),
        main_paragraph("[Replace with your keywords — 3–5 MeSH terms, semicolon-separated.]"),
        # Introduction starts on a new page (requirement 2)
        main_heading("Introduction", page_break_before=True),
        main_paragraph("[Clinical context, brief literature background, and why this case is novel or instructive (150–250 words).]"),
        {"kind": "intro_demo"},
        main_heading("Case presentation"),
        main_paragraph("[Patient demographics (age, sex), chief complaint, history of present illness, past medical/surgical history, physical examination findings, diagnostic workup (labs, imaging, pathology), treatment/surgical intervention with technique details, postoperative course, complications, follow-up, and outcome.]"),
        main_paragraph("[Replace with your case presentation content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Discussion"),
        main_paragraph("[Comparison with existing literature, pathophysiology review, differential diagnoses considered, why the chosen management was appropriate, and limitations of the report.]"),
        main_paragraph("[Replace with your discussion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Conclusion"),
        main_paragraph("[1–2 key clinical takeaways.]"),
        main_paragraph("[Replace with your conclusion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
    ]


def case_series_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        main_paragraph("[Replace with your structured abstract — maximum 350 words.]"),
        sub_heading("Background"),
        main_paragraph("[Replace with your background content for this section of the abstract.]"),
        sub_heading("Methods"),
        main_paragraph("[Replace with your methods content for this section of the abstract.]"),
        sub_heading("Results"),
        main_paragraph("[Replace with your results content for this section of the abstract.]"),
        sub_heading("Discussion"),
        main_paragraph("[Replace with your discussion content for this section of the abstract.]"),
        sub_heading("Conclusion"),
        main_paragraph("[Replace with your conclusion content for this section of the abstract.]"),
        main_heading("Keywords"),
        main_paragraph("[Replace with your keywords — 3–6 MeSH terms, semicolon-separated.]"),
        main_heading("Introduction", page_break_before=True),
        main_paragraph("[Clinical context, gap in literature, and the purpose of the series (200–300 words).]"),
        {"kind": "intro_demo"},
        main_heading("Methods"),
        main_paragraph("[Patient selection criteria (inclusion/exclusion), time period, data collection methods, outcome measures, and ethical approvals/consent. Reference the JBI critical appraisal checklist used.]"),
        main_paragraph("[Replace with your methods content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Results"),
        main_paragraph("[Patient demographics (table strongly recommended), case-by-case or aggregated findings, clinical outcomes, complications, and follow-up duration.]"),
        main_paragraph("[Replace with your results content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Discussion"),
        main_paragraph("[Synthesis of findings across cases, comparison with published literature, clinical implications, strengths, and limitations.]"),
        main_paragraph("[Replace with your discussion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Conclusion"),
        main_paragraph("[Summary of key findings and recommendations for practice.]"),
        main_paragraph("[Replace with your conclusion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
    ]


def surgical_technique_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        main_paragraph("[Replace with your unstructured abstract — maximum 200 words. No section headings; one continuous paragraph.]"),
        main_heading("Keywords"),
        main_paragraph("[Replace with your keywords — 3–5 MeSH terms, semicolon-separated.]"),
        main_heading("Introduction", page_break_before=True),
        main_paragraph("[Clinical problem addressed, limitations of current techniques, and rationale for the new or modified approach (150–200 words).]"),
        {"kind": "intro_demo"},
        main_heading("Surgical technique"),
        main_paragraph("[Step-by-step procedural description including patient positioning, approach, instruments required, key steps with intraoperative photos/diagrams, and closure. Video material is encouraged.]"),
        main_paragraph("[Replace with your surgical technique content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Discussion"),
        main_paragraph("[Advantages over existing techniques, potential limitations, learning curve, tips and pitfalls (200–400 words).]"),
        main_paragraph("[Replace with your discussion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Conclusion"),
        main_paragraph("[When to consider this technique.]"),
        main_paragraph("[Replace with your conclusion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
    ]


def images_in_orthopedics_sections() -> List[dict]:
    return [
        main_heading("Keywords"),
        main_paragraph("[Replace with your keywords — 3 keywords, semicolon-separated.]"),
        main_heading("Title", page_break_before=True),
        main_paragraph("[Descriptive, hints at the diagnosis or finding (max 15 words). The blinded manuscript file does not include the title — the title goes on the separate Title Page.]"),
        main_heading("Clinical description"),
        main_paragraph("[Brief clinical context including patient demographics, presenting complaint, key findings, and diagnosis (300–500 words). No formal section headings within this block.]"),
        {"kind": "intro_demo"},
        main_heading("Teaching point"),
        main_paragraph("[1–2 sentences on the clinical learning point.]"),
        main_paragraph("[Replace with your teaching point content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
    ]


def letter_to_editor_sections() -> List[dict]:
    return [
        main_heading("Body", page_break_before=True),
        main_paragraph("[No formal section headings required. Must reference the specific OSCRSJ article being discussed (by DOI). State the point of agreement or disagreement with supporting evidence (maximum 600 words).]"),
        {"kind": "intro_demo"},
        main_paragraph("[Replace with your body content. Letters to the Editor do not need separate Methods/Results/Discussion blocks — write as one continuous argument with paragraph breaks where natural.]"),
    ]


def review_article_sections() -> List[dict]:
    return [
        main_heading("Abstract"),
        main_paragraph("[Replace with your structured abstract — maximum 350 words.]"),
        sub_heading("Background"),
        main_paragraph("[Replace with your background content for this section of the abstract.]"),
        sub_heading("Methods"),
        main_paragraph("[Replace with your methods content for this section of the abstract.]"),
        sub_heading("Results"),
        main_paragraph("[Replace with your results content for this section of the abstract.]"),
        sub_heading("Conclusion"),
        main_paragraph("[Replace with your conclusion content for this section of the abstract.]"),
        main_heading("Keywords"),
        main_paragraph("[Replace with your keywords — 3–6 MeSH terms, semicolon-separated.]"),
        main_heading("Introduction", page_break_before=True),
        main_paragraph("[Clinical context and why this review is needed.]"),
        {"kind": "intro_demo"},
        main_heading("Methods"),
        main_paragraph("[Search strategy, databases searched, inclusion/exclusion criteria, and date range. PRISMA flow diagram recommended where applicable.]"),
        main_paragraph("[Replace with your methods content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Results"),
        main_paragraph("[Organized by theme or chronology. Tables strongly recommended for comparative summaries.]"),
        main_paragraph("[Replace with your results content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Discussion"),
        main_paragraph("[Synthesis of findings, clinical implications, and gaps in current knowledge.]"),
        main_paragraph("[Replace with your discussion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
        main_heading("Conclusion"),
        main_paragraph("[Key findings and future directions.]"),
        main_paragraph("[Replace with your conclusion content. Use Heading 2 for any subsections and Heading 3 for any sub-subsections. Do not use Heading 4 or deeper.]"),
    ]


TEMPLATES = [
    {
        "filename": "oscrsj-template-case-report.docx",
        "running_title_placeholder": "[Short running title — case report]",
        "sections": case_report_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (10–25 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
    {
        "filename": "oscrsj-template-case-series.docx",
        "running_title_placeholder": "[Short running title — case series]",
        "sections": case_series_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (15–40 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
    {
        "filename": "oscrsj-template-surgical-technique.docx",
        "running_title_placeholder": "[Short running title — surgical technique]",
        "sections": surgical_technique_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (8–15 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
    {
        "filename": "oscrsj-template-images-in-orthopedics.docx",
        "running_title_placeholder": "[Short running title — images in orthopedics]",
        "sections": images_in_orthopedics_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (maximum 5 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
    {
        "filename": "oscrsj-template-letter-to-editor.docx",
        "running_title_placeholder": "[Short running title — letter to the editor]",
        "sections": letter_to_editor_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (maximum 5 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
    {
        "filename": "oscrsj-template-review-article.docx",
        "running_title_placeholder": "[Short running title — review article]",
        "sections": review_article_sections,
        "references_intro": "[Number references manually as 1, 2, 3 in citation order. Use Vancouver style (20–60 references). One reference per paragraph. Do NOT use Word's auto-numbered list feature. Include the DOI URL inline at the end of each reference: https://doi.org/10.xxxx/xxxxx (no 'doi:' prefix). Three placeholder references with live bookmarks (REF_1, REF_2, REF_3) follow — replace with your own and add bookmarks REF_4, REF_5, … as needed.]",
    },
]


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------


def render_intro_demo(doc) -> None:
    """Emit a 1-paragraph instruction + 1-paragraph worked example showing
    superscript citations linked to the bookmarked references at the end of
    the manuscript.
    """
    add_paragraph(doc, CITATION_HOWTO, italic=True)
    # Worked example: build the paragraph by mixing plain runs with three
    # superscript hyperlinks targeting REF_1, REF_2, REF_3.
    add_paragraph_with_inline_superscripts(
        doc,
        [
            (
                "[Worked example sentence — replace with your own. Recent "
                "series have reported similar outcomes in this patient "
                "population",
                None,
            ),
            ("1", "REF_1"),
            (",", None),
            ("2", "REF_2"),
            (
                ", and a recent meta-analysis confirmed the trend",
                None,
            ),
            ("3", "REF_3"),
            (
                ". The superscript numbers above are live hyperlinks — "
                "Ctrl+Click (Cmd+Click on Mac) jumps to the matching "
                "Reference at the end of this manuscript.]",
                None,
            ),
        ],
    )


def render_template(spec: dict) -> Path:
    doc = Document()
    configure_document_defaults(doc)
    section = doc.sections[0]
    configure_section(
        section,
        running_title_placeholder=spec["running_title_placeholder"],
    )

    # ----- main manuscript body -----
    sections = spec["sections"]()
    is_first_block = True
    for entry in sections:
        kind = entry["kind"]
        if kind == "h1":
            # Empty line ABOVE every heading except the very first paragraph
            # of the document — that paragraph already starts on page 1.
            if not is_first_block:
                add_blank_paragraph(doc)
            add_heading_paragraph(
                doc,
                entry["text"],
                centered=False,
                page_break_before=entry.get("page_break_before", False),
            )
            is_first_block = False
        elif kind == "h2":
            add_blank_paragraph(doc)
            add_heading_paragraph(doc, entry["text"], centered=False)
            is_first_block = False
        elif kind == "para":
            add_paragraph(doc, entry["text"])
            is_first_block = False
        elif kind == "intro_demo":
            render_intro_demo(doc)
            is_first_block = False
        else:
            raise ValueError(f"Unknown section kind: {kind!r}")

    # ----- References (new page, centered heading, NO blank below) -----
    add_heading_paragraph(doc, "References", centered=True, page_break_before=True)
    add_paragraph(doc, spec["references_intro"], italic=True)
    # Three pre-seeded references with bookmarks REF_1 / REF_2 / REF_3.
    for n in (1, 2, 3):
        ref_para = add_paragraph(
            doc,
            f"{n}.  [Replace with your reference {n} in Vancouver/NLM style. "
            "Example: Surname AB, Surname CD, Surname EF. Title of article. "
            "J Abbreviated Name. YYYY;Vol(Issue):Pages. "
            "https://doi.org/10.xxxx/xxxxx]",
            indent_hanging_inches=0.25,
        )
        add_bookmark(ref_para, f"REF_{n}")

    # ----- Figure legends (new page, centered heading, NO blank below) -----
    add_heading_paragraph(doc, "Figure legends", centered=True, page_break_before=True)
    add_paragraph(
        doc,
        "[Provide one figure legend per figure here. Each legend begins with "
        "the figure number (e.g., 'Figure 1.') followed by a brief descriptive "
        "caption. Submit each figure as a separate high-resolution image file "
        "(TIFF, PNG, or JPEG, minimum 300 DPI) — do NOT embed images in this "
        "manuscript file. Indicate placement in the body of the manuscript "
        "with '[Insert Figure 1 here]' callouts at the relevant points in "
        "the text.]",
        italic=True,
    )
    add_paragraph(doc, "[Figure 1.  Replace with your figure 1 legend.]")
    add_paragraph(doc, "[Figure 2.  Replace with your figure 2 legend.]")

    out_path = DOWNLOADS_DIR / spec["filename"]
    doc.save(str(out_path))
    return out_path


def main() -> None:
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Writing templates to: {DOWNLOADS_DIR}\n")
    for spec in TEMPLATES:
        path = render_template(spec)
        size_kb = path.stat().st_size / 1024
        print(f"  ✓ {spec['filename']:<48} {size_kb:6.1f} KB")
    print("\nDone.")


if __name__ == "__main__":
    main()
